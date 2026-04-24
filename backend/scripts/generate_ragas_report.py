"""
generate_ragas_report.py — Construye docs/reporte-RAGAS.docx a partir del CSV
+ summary.json producidos por run_ragas_eval.py.

Uso:
    docker compose exec backend python scripts/generate_ragas_report.py \\
        --summary scripts/ragas_runs/20260424_0410_baseline.summary.json \\
        --csv     scripts/ragas_runs/20260424_0410_baseline.csv \\
        --out     /app/docs/reporte-RAGAS.docx

Requiere python-docx (ya en requirements.txt vía langchain / optional).
Instalar explícito si falta: pip install python-docx
"""

from __future__ import annotations

import argparse
import csv
import json
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit("Falta python-docx. Instalar: pip install python-docx")


# ----------------------------------------------------------------------------
# Helpers de estilo
# ----------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 11
FONT_SIZE_TITLE = 22
NAVY = RGBColor(0x14, 0x2D, 0x5E)


def set_run_font(run, size=FONT_SIZE_BODY, bold=False, color: RGBColor | None = None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    size = {1: 18, 2: 14, 3: 12}.get(level, FONT_SIZE_BODY)
    set_run_font(run, size=size, bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def body(doc: Document, text: str, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run(text)
    run.text = text
    set_run_font(run)


def kv_table(doc: Document, pairs: list[tuple[str, str]]):
    table = doc.add_table(rows=len(pairs), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"
    for i, (k, v) in enumerate(pairs):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.text = ""
        c1.text = ""
        r0 = c0.paragraphs[0].add_run(k)
        set_run_font(r0, bold=True)
        r1 = c1.paragraphs[0].add_run(v)
        set_run_font(r1)


# ----------------------------------------------------------------------------
# Secciones
# ----------------------------------------------------------------------------
def portada(doc: Document, summary: dict):
    # Título
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Reporte de Validación RAGAS")
    set_run_font(run, size=FONT_SIZE_TITLE, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pipeline RAG — Tutor IA Generativa")
    set_run_font(run, size=16, bold=False, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Curso Aplicaciones Móviles · IESTP República Federal de Alemania")
    set_run_font(run, size=12)

    for _ in range(6):
        doc.add_paragraph()

    meta = [
        ("Autor (tesista)", "Roger Alessandro Zavaleta Marcelo"),
        ("Asesora (USAT)", "Mg. Reyes Burgos, Karla"),
        ("Coordinador piloto (IESTP RFA)", "Téc. Xavier Benites Marín"),
        ("Institución", "Universidad Católica Santo Toribio de Mogrovejo"),
        ("Escuela", "Ingeniería de Sistemas y Computación"),
        ("Sprint", "Sprint 4 · CRISP-DM Evaluation"),
        ("Iteración", summary.get("label", "baseline")),
        ("Fecha de evaluación", summary.get("timestamp", datetime.now().strftime("%Y-%m-%d"))),
    ]
    kv_table(doc, meta)
    add_page_break(doc)


def indice(doc: Document):
    heading(doc, "Índice", level=1)
    items = [
        "1. Introducción",
        "2. Metodología",
        "   2.1 Golden set",
        "   2.2 Métricas RAGAS",
        "   2.3 Adaptación a Ollama local",
        "3. Configuración del pipeline evaluado",
        "4. Resultados globales",
        "5. Resultados por dimensión",
        "   5.1 Por módulo",
        "   5.2 Por tipo de pregunta",
        "   5.3 Por dificultad",
        "6. Casos límite",
        "7. Verificación de umbrales",
        "8. Conclusiones y siguientes pasos",
        "Anexo A · Tabla de resultados por pregunta",
    ]
    for i in items:
        body(doc, i)
    add_page_break(doc)


def seccion_intro(doc: Document, summary: dict):
    heading(doc, "1. Introducción", level=1)
    body(doc,
        "Este reporte documenta la validación del pipeline RAG del Sistema Tutor "
        "Inteligente (STI) construido como objeto de tesis, aplicando el framework "
        "RAGAS (Retrieval-Augmented Generation Assessment). La evaluación busca "
        "cumplir los umbrales establecidos en el Objetivo Específico 2 (OE2) "
        "y constituye el entregable principal del Sprint 4 (CRISP-DM Evaluation)."
    )
    body(doc,
        f"La ejecución cubrió {summary.get('n_questions', 'N')} preguntas del "
        "golden set, cuidando diversidad por módulo del curso (M1-M3), tipo de "
        "pregunta (conceptual, código, aplicación) y dificultad (easy/medium/hard). "
        "Todas las llamadas se realizaron con modelos auto-hospedados en Ollama, "
        "respetando la restricción de privacidad del proyecto (sin APIs pagas)."
    )


def seccion_metodologia(doc: Document, summary: dict):
    heading(doc, "2. Metodología", level=1)

    heading(doc, "2.1 Golden set", level=2)
    body(doc,
        "El golden set está compuesto por 30 preguntas distribuidas entre M1 "
        "(Fundamentos), M2 (Lógica Kotlin) y M3 (Interfaces UI). Cada pregunta "
        "incluye: enunciado, respuesta de referencia (ground_truth), palabras "
        "clave esperadas en el contexto, tipo (conceptual/code/application) "
        "y dificultad. El archivo fuente es "
        "backend/tests/fixtures/golden_set.json, versionado en el repositorio."
    )

    heading(doc, "2.2 Métricas RAGAS", level=2)
    body(doc,
        "Se implementaron las cuatro métricas canónicas de RAGAS, adaptadas para "
        "operar con un LLM juez local. Todas son calculadas por respuesta y "
        "luego promediadas globalmente y por dimensión:"
    )
    bullet(doc,
        "Faithfulness: proporción de claims extraídos de la respuesta que están "
        "respaldados por el contexto recuperado. Umbral objetivo ≥ 0,75."
    )
    bullet(doc,
        "Answer relevancy: similitud semántica (coseno con mxbai-embed-large) "
        "entre la pregunta original y preguntas sintéticas generadas a partir "
        "de la respuesta. Umbral objetivo ≥ 0,70."
    )
    bullet(doc,
        "Context precision: precisión promedio @k ponderada por posición — "
        "proporción de chunks útiles en el top-k recuperado."
    )
    bullet(doc,
        "Context recall: proporción de sentencias del ground_truth que son "
        "atribuibles al contexto recuperado."
    )

    heading(doc, "2.3 Adaptación a Ollama local", level=2)
    body(doc,
        "El paquete ragas estándar asume OpenAI por defecto. Dado que el proyecto "
        "exige privacidad total, se implementó una versión equivalente de las "
        "cuatro métricas invocando a Ollama (qwen2.5:7b-instruct-q4_K_M como "
        "juez y generador, y mxbai-embed-large para similitud). El juez usa "
        "temperatura 0 para determinismo; el generador usa 0,7 para diversidad "
        "en answer_relevancy. Las respuestas del LLM se solicitan en formato "
        "JSON (format=\"json\") y el parser tolera estructuras con objetos o "
        "booleanos desnudos. La evaluación es robusta ante fallos puntuales "
        "del LLM: si una métrica retorna None, la pregunta se preserva con las "
        "métricas exitosas."
    )


def seccion_config(doc: Document, summary: dict):
    heading(doc, "3. Configuración del pipeline evaluado", level=1)
    cfg = summary.get("config", {})
    rows = [
        ("Modelo LLM", cfg.get("llm", "-")),
        ("Modelo embeddings", cfg.get("embed", "-")),
        ("Chunk size", str(cfg.get("chunk_size", "-"))),
        ("Chunk overlap", str(cfg.get("chunk_overlap", "-"))),
        ("Top-k", str(cfg.get("top_k", "-"))),
        ("Threshold cosine", str(cfg.get("threshold", "-"))),
        ("Iteración", summary.get("label", "baseline")),
    ]
    kv_table(doc, rows)


def seccion_globales(doc: Document, summary: dict):
    heading(doc, "4. Resultados globales", level=1)
    glob = summary.get("global", {})
    rows = [
        ("Faithfulness", f"{glob.get('faithfulness')}"),
        ("Answer relevancy", f"{glob.get('answer_relevancy')}"),
        ("Context precision", f"{glob.get('context_precision')}"),
        ("Context recall", f"{glob.get('context_recall')}"),
    ]
    kv_table(doc, rows)

    body(doc, "")
    body(doc,
        f"La evaluación completa cubrió {summary.get('n_questions', 'N')} "
        "preguntas. Las respuestas fueron generadas fresh (sin caché Redis) "
        "para garantizar que las métricas reflejen el comportamiento real del "
        "pipeline tras la ingesta del corpus."
    )


def _metric_table(doc: Document, title: str, data: dict):
    heading(doc, title, level=2)
    if not data:
        body(doc, "Sin datos.")
        return

    cols = ["", "n", "faithfulness", "answer_relevancy", "context_precision", "context_recall"]
    table = doc.add_table(rows=1 + len(data), cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"

    for i, c in enumerate(cols):
        cell = table.cell(0, i)
        cell.text = ""
        r = cell.paragraphs[0].add_run(c.replace("_", " ").capitalize() if c else "Dimensión")
        set_run_font(r, bold=True)

    for row, (k, v) in enumerate(sorted(data.items()), start=1):
        table.cell(row, 0).text = str(k)
        table.cell(row, 1).text = str(v.get("n", ""))
        table.cell(row, 2).text = f"{v.get('faithfulness')}"
        table.cell(row, 3).text = f"{v.get('answer_relevancy')}"
        table.cell(row, 4).text = f"{v.get('context_precision')}"
        table.cell(row, 5).text = f"{v.get('context_recall')}"
        for c in table.rows[row].cells:
            for p in c.paragraphs:
                for r in p.runs:
                    set_run_font(r)


def seccion_dimensiones(doc: Document, summary: dict):
    heading(doc, "5. Resultados por dimensión", level=1)
    _metric_table(doc, "5.1 Por módulo", summary.get("by_module", {}))
    _metric_table(doc, "5.2 Por tipo de pregunta", summary.get("by_type", {}))
    _metric_table(doc, "5.3 Por dificultad", summary.get("by_difficulty", {}))


def seccion_casos_limite(doc: Document, rows: list[dict]):
    heading(doc, "6. Casos límite", level=1)
    body(doc,
        "A continuación se listan las cinco preguntas con menor faithfulness; "
        "estos casos merecen análisis individual en iteraciones futuras."
    )

    def safe_float(v):
        try:
            return float(v)
        except (TypeError, ValueError):
            return 1.0

    worst = sorted(rows, key=lambda r: safe_float(r.get("faithfulness", 1.0)))[:5]

    table = doc.add_table(rows=1 + len(worst), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"
    for i, h in enumerate(["Q", "Módulo", "Tipo", "Faithfulness", "Pregunta"]):
        cell = table.cell(0, i)
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, bold=True)
    for r, w in enumerate(worst, start=1):
        table.cell(r, 0).text = str(w.get("id", ""))
        table.cell(r, 1).text = str(w.get("module", ""))
        table.cell(r, 2).text = str(w.get("type", ""))
        table.cell(r, 3).text = str(w.get("faithfulness", ""))
        table.cell(r, 4).text = str(w.get("question", ""))[:120]
        for c in table.rows[r].cells:
            for p in c.paragraphs:
                for rr in p.runs:
                    set_run_font(rr)


def seccion_umbrales(doc: Document, summary: dict):
    heading(doc, "7. Verificación de umbrales", level=1)
    g = summary.get("global", {})
    passed = summary.get("pass", {})

    rows = [
        ("Faithfulness", f"{g.get('faithfulness')}", "≥ 0,75", "✓" if passed.get("faithfulness") else "✗"),
        ("Answer relevancy", f"{g.get('answer_relevancy')}", "≥ 0,70", "✓" if passed.get("answer_relevancy") else "✗"),
    ]

    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"
    for i, h in enumerate(["Métrica", "Resultado", "Umbral objetivo", "Estado"]):
        cell = table.cell(0, i)
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.cell(r_idx, c_idx).text = str(val)
            for p in table.cell(r_idx, c_idx).paragraphs:
                for rr in p.runs:
                    set_run_font(rr)


def seccion_conclusiones(doc: Document, summary: dict):
    heading(doc, "8. Conclusiones y siguientes pasos", level=1)

    g = summary.get("global", {})
    passed = summary.get("pass", {})
    all_pass = all(passed.values()) if passed else False

    if all_pass:
        body(doc,
            "El pipeline RAG del STI cumple los umbrales RAGAS definidos en el "
            "OE2: faithfulness ≥ 0,75 y answer_relevancy ≥ 0,70. El sistema "
            "produce respuestas fieles al corpus y pertinentes al enunciado "
            "del estudiante, condición habilitante para avanzar al piloto con "
            "10-15 estudiantes del IESTP RFA (Sprint 8)."
        )
    else:
        body(doc,
            "El pipeline en la iteración actual NO alcanza todos los umbrales RAGAS "
            "definidos. Se proponen las siguientes iteraciones antes del siguiente "
            "entregable:"
        )

    bullet(doc,
        "Iteración v2: comparar chunking recursivo (500/50) contra chunking "
        "semántico con solapamiento del 15 % (configuración v4.0)."
    )
    bullet(doc,
        "Iteración v3: ajustar threshold de similitud coseno (0,70 vs 0,65) "
        "manteniendo top-k=5."
    )
    bullet(doc,
        "Analizar los casos con menor faithfulness (Sección 6) para detectar "
        "temas sub-representados en el corpus ingestado y reforzarlos."
    )
    bullet(doc,
        "Ampliar el golden set a M4 y M5 en el Sprint 6 para cobertura "
        "diagnóstica completa."
    )

    body(doc, "")
    body(doc,
        "Todos los scripts de generación y el CSV detallado se conservan bajo "
        "backend/scripts/ragas_runs/ como insumo reproducible para la discusión "
        "del Informe Final (Sprint 8) y la matriz de trazabilidad ISO/IEC 25010 "
        "(Sprint 7)."
    )


def anexo_tabla(doc: Document, rows: list[dict]):
    add_page_break(doc)
    heading(doc, "Anexo A · Tabla de resultados por pregunta", level=1)

    cols = ["Q", "Mód.", "Tipo", "Dif.", "F", "AR", "CP", "CR", "s"]
    table = doc.add_table(rows=1 + len(rows), cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light List Accent 1"
    for i, h in enumerate(cols):
        cell = table.cell(0, i)
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_font(r, bold=True, size=9)

    for i, w in enumerate(rows, start=1):
        vals = [
            w.get("id", ""),
            w.get("module", ""),
            (w.get("type", "") or "")[:5],
            (w.get("difficulty", "") or "")[:4],
            w.get("faithfulness", ""),
            w.get("answer_relevancy", ""),
            w.get("context_precision", ""),
            w.get("context_recall", ""),
            w.get("elapsed_s", ""),
        ]
        for ci, v in enumerate(vals):
            table.cell(i, ci).text = str(v)
            for p in table.cell(i, ci).paragraphs:
                for rr in p.runs:
                    set_run_font(rr, size=9)


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------
def main(summary_path: Path, csv_path: Path, out_path: Path):
    summary = json.loads(summary_path.read_text(encoding="utf-8"))

    rows = []
    with open(csv_path, encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            # Coerce numeric
            for k in ("faithfulness", "answer_relevancy", "context_precision", "context_recall", "elapsed_s"):
                try:
                    r[k] = float(r[k]) if r.get(k) else None
                except ValueError:
                    r[k] = None
            rows.append(r)

    doc = Document()

    # Default style → Times New Roman 11
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_BODY)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # Márgenes 2.5 cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    portada(doc, summary)
    indice(doc)
    seccion_intro(doc, summary)
    seccion_metodologia(doc, summary)
    seccion_config(doc, summary)
    seccion_globales(doc, summary)
    seccion_dimensiones(doc, summary)
    seccion_casos_limite(doc, rows)
    seccion_umbrales(doc, summary)
    seccion_conclusiones(doc, summary)
    anexo_tabla(doc, rows)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"✓ Generado: {out_path}  ({out_path.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--summary", required=True)
    parser.add_argument("--csv", required=True)
    parser.add_argument("--out", default="/app/docs/reporte-RAGAS.docx")
    args = parser.parse_args()
    main(Path(args.summary), Path(args.csv), Path(args.out))
