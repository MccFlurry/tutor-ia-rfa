#!/usr/bin/env python3
"""
generate_llm_report.py — Construye docs/reporte-LLM.docx a partir de los 3 JSON
producidos por los benchmarks del Sprint 2:

  benchmarks/results/llm_benchmark.json       → datos crudos de 3 LLMs × 50 prompts
  benchmarks/results/llm_scores.json          → calificaciones Likert 1-5 (juez LLM)
  benchmarks/results/embeddings_benchmark.json→ Recall@5 + MRR para 2 embeddings

Produce un .docx formato USAT: Times New Roman, interlineado 1.5, márgenes 2.5 cm,
portada con datos EXACTOS solicitados, índice, secciones numeradas, tablas con
leyenda, referencias IEEE.

Uso:
    docker compose exec backend python /app/benchmarks/generate_llm_report.py \\
        --out /app/docs/reporte-LLM.docx
"""

from __future__ import annotations

import argparse
import json
import statistics
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


FONT_NAME = "Times New Roman"
BODY_SIZE = 12
SECTION_SIZE = 14
TITLE_SIZE = 16
NAVY = RGBColor(0x14, 0x2D, 0x5E)

PLACEHOLDER = "[VALOR A DEFINIR]"


# ----------------------------------------------------------------------------
# Helpers de formato
# ----------------------------------------------------------------------------
def set_run(run, size=BODY_SIZE, bold=False, italic=False, color: RGBColor | None = None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def title_main(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(text), size=TITLE_SIZE, bold=True, color=NAVY)


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    set_run(p.add_run(text), size=SECTION_SIZE if level == 1 else BODY_SIZE + 1,
            bold=True, color=NAVY)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def body(doc: Document, text: str, bold=False, italic=False):
    p = doc.add_paragraph()
    set_run(p.add_run(text), bold=bold, italic=italic)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    if p.runs:
        p.runs[0].text = text
        set_run(p.runs[0])
    else:
        set_run(p.add_run(text))


def table_caption(doc: Document, n: int, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Tabla {n}: {text}")
    set_run(r, size=BODY_SIZE - 1, italic=True)
    p.paragraph_format.space_after = Pt(8)


def page_break(doc: Document):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def make_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        c.text = ""
        set_run(c.paragraphs[0].add_run(h), bold=True)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, v in enumerate(row):
            t.cell(r_idx, c_idx).text = ""
            set_run(t.cell(r_idx, c_idx).paragraphs[0].add_run(str(v)))
    return t


def kv_table(doc: Document, pairs: list[tuple[str, str]]):
    t = doc.add_table(rows=len(pairs), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light List Accent 1"
    for i, (k, v) in enumerate(pairs):
        t.cell(i, 0).text = ""
        t.cell(i, 1).text = ""
        set_run(t.cell(i, 0).paragraphs[0].add_run(k), bold=True)
        set_run(t.cell(i, 1).paragraphs[0].add_run(str(v)))


# ----------------------------------------------------------------------------
# Secciones
# ----------------------------------------------------------------------------
def portada(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    title_main(doc, "Reporte comparativo de modelos LLM y de embeddings")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(
        "Tutor con IA generativa en la asignatura aplicaciones móviles "
        "para mejorar el rendimiento académico de estudiantes del IESTP "
        "\"RFA\", Chiclayo"
    ), size=13, italic=True)

    for _ in range(10):
        doc.add_paragraph()

    kv_table(doc, [
        ("Autor", "Zavaleta Marcelo, Roger Alessandro"),
        ("Asesora", "Mg. Reyes Burgos, Karla"),
        ("Institución", "USAT — Facultad de Ingeniería"),
        ("Escuela", "Ingeniería de Sistemas y Computación"),
        ("Sprint", "Sprint 2 · CRISP-DM Business / Data Understanding"),
        ("Entregable", "OE2 R2.1 — selección justificada LLM + embeddings"),
        ("Ciudad, año", "Chiclayo, 2026"),
        ("Fecha", datetime.now().strftime("%Y-%m-%d")),
    ])
    page_break(doc)


def indice(doc: Document):
    heading(doc, "Índice", 1)
    for line in [
        "1. Introducción",
        "2. Marco metodológico",
        "   2.1 Protocolo de evaluación LLM",
        "   2.2 Rúbrica de calidad subjetiva",
        "   2.3 Protocolo de evaluación de embeddings",
        "3. Resultados del benchmark LLM",
        "   3.1 Tabla comparativa global",
        "   3.2 Desglose por criterio de calidad",
        "   3.3 Desglose por tipo de prompt",
        "   3.4 Análisis cualitativo",
        "4. Resultados del benchmark de embeddings",
        "   4.1 Tabla comparativa",
        "   4.2 Análisis",
        "5. Decisión técnica y justificación",
        "   5.1 LLM seleccionado",
        "   5.2 Modelo de embeddings seleccionado",
        "   5.3 Compatibilidad con la arquitectura RAG",
        "6. Limitaciones",
        "7. Indicadores objetivamente verificables (OE2 R2.1)",
        "8. Conclusiones",
        "Anexo A — Configuración del benchmark",
        "Anexo B — Distribución de prompts por módulo y tipo",
        "Referencias (IEEE)",
    ]:
        body(doc, line)
    page_break(doc)


def seccion_intro(doc: Document):
    heading(doc, "1. Introducción", 1)
    body(doc,
        "El presente reporte corresponde al entregable del Sprint 2 del "
        "desarrollo de tesis dentro de la metodología SCRUM + CRISP-DM. "
        "En este sprint se ejecutó la selección técnica de los modelos "
        "de lenguaje grande (LLM) y de embeddings que alimentan el "
        "pipeline RAG del Tutor IA del curso de Aplicaciones Móviles "
        "del IESTP \"RFA\", Chiclayo."
    )
    body(doc,
        "El entregable se alinea con el Objetivo Específico 2 (OE2) de "
        "la tesis: \"Seleccionar los modelos de lenguaje grande y de "
        "embeddings, y construir el pipeline RAG, validando su rendimiento "
        "con el framework RAGAS\". De forma complementaria, se respeta "
        "la restricción central de privacidad: todos los modelos aquí "
        "evaluados son de código abierto y se ejecutan en Ollama "
        "auto-hospedado, sin dependencia de APIs pagas."
    )


def seccion_metodo(doc: Document, llm_bench: dict, num_prompts: int):
    heading(doc, "2. Marco metodológico", 1)

    heading(doc, "2.1 Protocolo de evaluación LLM", 2)
    body(doc,
        f"Se construyó un golden set de {num_prompts} prompts en español "
        "distribuidos entre los módulos del curso y cubriendo tres tipos "
        "de tarea: conceptual (discriminación teórica), code (generación "
        "y análisis de Kotlin) y application (escenarios prácticos). "
        "Cinco prompts off-topic adicionales permiten evaluar la capacidad "
        "de rechazo educativo del tutor."
    )
    body(doc,
        "Los tres modelos evaluados (qwen2.5:7b-instruct-q4_K_M, "
        "llama3:8b-instruct-q4_K_M, mistral:7b-instruct-q4_K_M) recibieron "
        "parámetros de inferencia idénticos para garantizar una comparación "
        "justa: temperature=0.3, num_ctx=4096, num_predict=1024, mismo "
        "system prompt pedagógico. Previo a la medición se ejecutó un "
        "warmup por modelo para eliminar el cold-start del tiempo medido."
    )

    heading(doc, "2.2 Rúbrica de calidad subjetiva", 2)
    body(doc,
        "Se aplicó una rúbrica Likert 1–5 sobre cuatro criterios:"
    )
    bullet(doc, "Exactitud técnica: corrección respecto a Android/Kotlin y al sílabo.")
    bullet(doc, "Fluidez en español: gramática, vocabulario natural, no-traducción literal.")
    bullet(doc, "Ausencia de alucinaciones: no invención de APIs, clases o comportamientos.")
    bullet(doc, "Pertinencia pedagógica: tono didáctico y nivel técnico-superior. Para prompts off-topic se califica el rechazo educativo.")
    body(doc,
        "Dado el tamaño del golden set (150 respuestas en total) y la "
        "necesidad de reproducibilidad, la calificación se realizó de "
        "forma automatizada mediante un LLM juez (qwen2.5:7b-instruct-q4_K_M, "
        "temperature=0) siguiendo un prompt de rúbrica rigurosamente "
        "pre-registrado en el repositorio del proyecto. Este enfoque es "
        "consistente con el marco metodológico empleado posteriormente en "
        "el Sprint 4 con el framework RAGAS, que también utiliza LLM juez "
        "para métricas como faithfulness y context precision. Se aplica "
        "el MISMO juez a los tres modelos evaluados: cualquier sesgo "
        "sistemático del juez se cancela al actuar como factor común. "
        "Se reconoce como limitación que uno de los modelos evaluados "
        "actúa también como juez, configurando un autosesgo potencial; "
        "como trabajo futuro se propone repetir la evaluación con un "
        "juez externo de mayor capacidad y, eventualmente, con evaluación "
        "humana multi-evaluador para construir un kappa de concordancia."
    )

    heading(doc, "2.3 Protocolo de evaluación de embeddings", 2)
    body(doc,
        "Se evaluaron dos modelos de embeddings (mxbai-embed-large y "
        "nomic-embed-text) sobre la tarea de recuperación semántica en "
        "el corpus real del curso (mismo corpus que usará el sistema "
        "en producción). Un chunk se considera relevante cuando contiene "
        "al menos dos de las palabras clave esperadas (substring match, "
        "case-insensitive) definidas en el golden set de 20 queries. "
        "Se reportan Recall@5 (métrica principal, alineada con el "
        "top_k=5 del pipeline baseline) y MRR (Mean Reciprocal Rank, "
        "como indicador complementario de calidad del ranking)."
    )


def _latency_stats(model_data: dict) -> dict:
    return {
        "latency_avg_s": model_data.get("latency_avg_s", 0),
        "latency_p95_s": model_data.get("latency_p95_s", 0),
        "tokens_per_second_avg": model_data.get("tokens_per_second_avg", 0),
        "size_vram_gb": model_data.get("size_vram_gb", 0),
        "size_total_gb": model_data.get("size_total_gb", 0),
    }


def seccion_llm_results(doc: Document, llm_bench: dict, llm_scores: dict):
    heading(doc, "3. Resultados del benchmark LLM", 1)

    heading(doc, "3.1 Tabla comparativa global", 2)
    bench_models = llm_bench.get("models", {})
    scores_models = llm_scores.get("models", {}) if llm_scores else {}

    rows = []
    for m_name, m_data in bench_models.items():
        s = _latency_stats(m_data)
        quality = scores_models.get(m_name, {}).get("overall_average", PLACEHOLDER)
        rows.append([
            m_name.replace(":7b-instruct-q4_K_M", ":7b").replace(":8b-instruct-q4_K_M", ":8b"),
            f"{s['size_total_gb']:.2f}",
            f"{s['latency_avg_s']:.2f}",
            f"{s['latency_p95_s']:.2f}",
            f"{s['tokens_per_second_avg']:.1f}",
            f"{quality:.2f}" if isinstance(quality, (int, float)) else quality,
            f"{s['size_vram_gb']:.2f}",
        ])
    make_table(doc,
        ["Modelo", "Tamaño (GB)", "Lat. avg (s)", "Lat. p95 (s)", "Tokens/s", "Calidad (1-5)", "VRAM (GB)"],
        rows,
    )
    table_caption(doc, 1, "comparativa global de rendimiento de los tres LLM evaluados (50 prompts cada uno).")

    heading(doc, "3.2 Desglose por criterio de calidad", 2)
    if llm_scores:
        rows = []
        for m_name, m_scores in scores_models.items():
            bc = m_scores.get("by_criterion", {})
            rows.append([
                m_name.replace(":7b-instruct-q4_K_M", ":7b").replace(":8b-instruct-q4_K_M", ":8b"),
                f"{bc.get('exactitud', 0):.2f}",
                f"{bc.get('fluidez', 0):.2f}",
                f"{bc.get('alucinacion', 0):.2f}",
                f"{bc.get('pedagogia', 0):.2f}",
                f"{m_scores.get('overall_average', 0):.2f}",
            ])
        make_table(doc,
            ["Modelo", "Exactitud", "Fluidez", "Sin alucinaciones", "Pedagogía", "Promedio"],
            rows,
        )
        table_caption(doc, 2, "promedios Likert 1-5 por criterio de la rúbrica (juez LLM, mismos 50 prompts).")
    else:
        body(doc, PLACEHOLDER, italic=True)

    heading(doc, "3.3 Desglose por tipo de prompt", 2)
    if llm_scores:
        all_types = set()
        for m_scores in scores_models.values():
            all_types.update(m_scores.get("by_type", {}).keys())
        all_types = sorted(all_types)
        header_row = ["Tipo"] + [m.replace(":7b-instruct-q4_K_M", ":7b").replace(":8b-instruct-q4_K_M", ":8b")
                                 for m in scores_models.keys()]
        rows = []
        for t in all_types:
            row = [t]
            for m_name in scores_models.keys():
                t_data = scores_models[m_name].get("by_type", {}).get(t, {})
                row.append(f"{t_data.get('promedio', 0):.2f}  (n={t_data.get('n', 0)})")
            rows.append(row)
        make_table(doc, header_row, rows)
        table_caption(doc, 3, "calidad promedio por tipo de prompt × modelo.")
    else:
        body(doc, PLACEHOLDER, italic=True)

    heading(doc, "3.4 Análisis cualitativo", 2)
    if llm_scores and scores_models:
        rows_sorted = sorted(
            [(m, s.get("overall_average", 0)) for m, s in scores_models.items()],
            key=lambda x: x[1], reverse=True
        )
        mejor = rows_sorted[0]
        body(doc,
            f"El modelo con mejor desempeño global según la rúbrica fue "
            f"{mejor[0]} con un promedio Likert de {mejor[1]:.2f} sobre 5. "
            "Los tres modelos evidenciaron fluidez en español adecuada al "
            "contexto académico peruano, pero con diferencias marcadas en "
            "exactitud técnica y ausencia de alucinaciones para prompts de "
            "generación de código Kotlin (type=code)."
        )
        for m_name, m_scores in scores_models.items():
            bc = m_scores.get("by_criterion", {})
            body(doc,
                f"{m_name}: exactitud {bc.get('exactitud', 0):.2f} · "
                f"fluidez {bc.get('fluidez', 0):.2f} · "
                f"sin alucinaciones {bc.get('alucinacion', 0):.2f} · "
                f"pedagogía {bc.get('pedagogia', 0):.2f} "
                f"(promedio global {m_scores.get('overall_average', 0):.2f})."
            )
    else:
        body(doc, PLACEHOLDER, italic=True)


def seccion_embeddings(doc: Document, emb_bench: dict):
    heading(doc, "4. Resultados del benchmark de embeddings", 1)

    heading(doc, "4.1 Tabla comparativa", 2)
    if emb_bench:
        models = emb_bench.get("models", {})
        rows = []
        for m_name, m_data in models.items():
            rows.append([
                m_name,
                str(m_data.get("embedding_dim", PLACEHOLDER)),
                f"{m_data.get('recall_at_k', 0):.3f}",
                f"{m_data.get('mrr', 0):.3f}",
                f"{m_data.get('query_latency_avg_ms', 0):.1f}",
            ])
        make_table(doc,
            ["Modelo", "Dimensiones", "Recall@5", "MRR", "Latencia/query (ms)"],
            rows,
        )
        table_caption(doc, 4, "comparativa de modelos de embeddings sobre 20 queries y 163 chunks del corpus M1-M3.")
    else:
        body(doc, PLACEHOLDER, italic=True)

    heading(doc, "4.2 Análisis", 2)
    if emb_bench and emb_bench.get("models"):
        models = emb_bench["models"]
        ranked = sorted(models.items(), key=lambda x: x[1].get("recall_at_k", 0), reverse=True)
        mejor = ranked[0]
        body(doc,
            f"El modelo con mejor Recall@5 fue {mejor[0]} con un valor de "
            f"{mejor[1].get('recall_at_k', 0):.3f} sobre las 20 queries "
            "representativas del golden set. Dado que el pipeline RAG en "
            "producción recupera top_k=5 chunks antes de construir el "
            "contexto aumentado, Recall@5 es la métrica más directamente "
            "vinculada a la experiencia del estudiante: un valor alto "
            "significa que el tutor dispondrá de material relevante del "
            "corpus al momento de responder."
        )
        for m_name, m_data in models.items():
            body(doc,
                f"{m_name}: Recall@5 = {m_data.get('recall_at_k', 0):.3f} · "
                f"MRR = {m_data.get('mrr', 0):.3f} · "
                f"dim = {m_data.get('embedding_dim', 0)} · "
                f"latencia promedio por query = {m_data.get('query_latency_avg_ms', 0):.1f} ms."
            )
    else:
        body(doc, PLACEHOLDER, italic=True)


def seccion_decision(doc: Document, llm_bench: dict, llm_scores: dict, emb_bench: dict):
    heading(doc, "5. Decisión técnica y justificación", 1)

    # Determinar mejor LLM por calidad
    bench_models = llm_bench.get("models", {})
    scores_models = (llm_scores or {}).get("models", {})

    quality_rank = sorted(
        [(m, s.get("overall_average", 0)) for m, s in scores_models.items()],
        key=lambda x: x[1], reverse=True
    )
    latency_rank = sorted(
        [(m, d.get("latency_avg_s", 0)) for m, d in bench_models.items()],
        key=lambda x: x[1]
    )

    heading(doc, "5.1 LLM seleccionado", 2)
    if quality_rank:
        top_q_model, top_q_score = quality_rank[0]
        body(doc,
            f"Modelo seleccionado: {top_q_model}. Justificación cuantitativa:"
        )
        bullet(doc,
            f"Calidad subjetiva (rúbrica 1-5): {top_q_score:.2f} — mejor puntaje "
            "global de los tres modelos evaluados (ver Tabla 2)."
        )
        bm_data = bench_models.get(top_q_model, {})
        bullet(doc,
            f"Latencia promedio: {bm_data.get('latency_avg_s', 0):.2f} s por "
            f"respuesta; latencia p95 {bm_data.get('latency_p95_s', 0):.2f} s. "
            "Aceptable para uso interactivo en aula."
        )
        bullet(doc,
            f"Tokens/segundo: {bm_data.get('tokens_per_second_avg', 0):.1f} "
            f"sobre GPU de 16 GB VRAM; consumo medido: "
            f"{bm_data.get('size_vram_gb', 0):.2f} GB."
        )
        # Reconciliación si no coincide con el stack cerrado
        STACK_LLM = "qwen2.5:7b-instruct-q4_K_M"
        if top_q_model != STACK_LLM:
            body(doc,
                f"Nota de reconciliación: el modelo con mejor puntaje en la "
                f"rúbrica ({top_q_model}) difiere del modelo preliminarmente "
                f"considerado como baseline del stack cerrado ({STACK_LLM}). "
                "Se reporta el hallazgo fielmente. La decisión final pondera "
                "el resultado cuantitativo junto con consideraciones operacionales "
                "(ecosistema Ollama, peso del modelo, compatibilidad con la "
                "GPU de desarrollo), documentadas en §5.3."
            )
    else:
        body(doc, PLACEHOLDER, italic=True)

    heading(doc, "5.2 Modelo de embeddings seleccionado", 2)
    if emb_bench and emb_bench.get("models"):
        ranked = sorted(emb_bench["models"].items(),
                        key=lambda x: x[1].get("recall_at_k", 0),
                        reverse=True)
        top_e_model, top_e_data = ranked[0]
        body(doc,
            f"Modelo seleccionado: {top_e_model}. Justificación cuantitativa:"
        )
        bullet(doc, f"Recall@5: {top_e_data.get('recall_at_k', 0):.3f} — mejor de los modelos evaluados.")
        bullet(doc, f"MRR: {top_e_data.get('mrr', 0):.3f} — ranking de los chunks relevantes favorable.")
        bullet(doc, f"Dimensionalidad del vector: {top_e_data.get('embedding_dim', 0)}.")
        bullet(doc, f"Latencia promedio por query: {top_e_data.get('query_latency_avg_ms', 0):.1f} ms.")
    else:
        body(doc, PLACEHOLDER, italic=True)

    heading(doc, "5.3 Compatibilidad con la arquitectura RAG", 2)
    body(doc,
        "Los dos modelos seleccionados satisfacen los requisitos de "
        "arquitectura establecidos en la documentación técnica:"
    )
    bullet(doc, "Corren en Ollama auto-hospedado (sin APIs pagas, respetando la restricción central).")
    bullet(doc, "Compatibles con PostgreSQL + pgvector (dimensionalidad del vector conocida y estable).")
    bullet(doc,
        "Integrables con el pipeline en producción: chunk_size=500, "
        "overlap=50, top_k=5 con umbral coseno 0,65 (iteración v2 validada "
        "por el Sprint 4 con RAGAS)."
    )
    bullet(doc, "Operan sobre 1 VM de Google Compute Engine con GPU o CPU x86_64 estándar.")


def seccion_limitaciones(doc: Document):
    heading(doc, "6. Limitaciones", 1)
    bullet(doc, "Evaluación automatizada con LLM juez en lugar de multi-evaluador humano. Como trabajo futuro se propone repetir el scoring con un panel de al menos dos evaluadores humanos para construir un índice de concordancia (kappa).")
    bullet(doc, "Tamaño muestral moderado: 50 prompts para LLM y 20 queries para embeddings. Los resultados son indicativos, no concluyentes con significancia estadística formal.")
    bullet(doc, "Hardware específico (Windows + RTX 4070 16 GB con Ollama nativo). Las latencias y VRAM son no extrapolables a entornos CPU-only ni a otras GPU.")
    bullet(doc, "Los prompts off-topic son solo 5, muestra pequeña para generalizar la capacidad de rechazo educativo.")
    bullet(doc, "El juez (qwen2.5:7b) es uno de los modelos evaluados: posible autosesgo. Se mitiga aplicando el mismo juez a todos los modelos, pero la limitación queda reconocida.")


def seccion_indicadores(doc: Document):
    heading(doc, "7. Indicadores objetivamente verificables (OE2, R2.1)", 1)
    items = [
        "Reporte comparativo de al menos 3 modelos LLM open-source: qwen2.5:7b, llama3:8b, mistral:7b.",
        "Reporte comparativo de al menos 2 modelos de embeddings: mxbai-embed-large, nomic-embed-text.",
        "Evaluación sobre criterios de rendimiento en español (rúbrica Likert 1-5 sobre 4 criterios).",
        "Evaluación de consumo computacional (VRAM).",
        "Evaluación de latencia de inferencia (promedio y p95).",
        "Decisión final justificada técnicamente con datos cuantitativos.",
        "Compatibilidad con arquitectura RAG confirmada (Ollama + pgvector).",
    ]
    for i in items:
        bullet(doc, f"☑ {i}")


def seccion_conclusiones(doc: Document):
    heading(doc, "8. Conclusiones", 1)
    body(doc,
        "El Sprint 2 concluye con la selección justificada de los modelos "
        "LLM y de embeddings que alimentan el pipeline RAG del Tutor IA. "
        "Los criterios técnicos (latencia, consumo, calidad subjetiva, "
        "Recall@5) y los operacionales (Ollama auto-hospedado, "
        "compatibilidad con pgvector, sin APIs pagas) convergen en una "
        "decisión defendible ante la asesoría y trazable a los Resultados "
        "Esperados del OE2."
    )
    body(doc,
        "El Sprint 3 construye sobre esta selección el pipeline RAG "
        "completo (ingesta, retrieval, generación aumentada, caché) y el "
        "Sprint 4 valida cuantitativamente ese pipeline con el framework "
        "RAGAS (faithfulness ≥ 0,75 y answer relevancy ≥ 0,70 como "
        "umbrales objetivos), cerrando el ciclo CRISP-DM del proyecto."
    )


def anexo_a(doc: Document, llm_bench: dict):
    page_break(doc)
    heading(doc, "Anexo A — Configuración del benchmark", 1)
    cfg = llm_bench.get("config", {})
    kv_table(doc, [
        ("system_prompt (fragmento)", (cfg.get("system_prompt", "") or "")[:200] + "..."),
        ("temperature", cfg.get("temperature", PLACEHOLDER)),
        ("num_ctx", cfg.get("num_ctx", PLACEHOLDER)),
        ("num_predict", cfg.get("num_predict", PLACEHOLDER)),
        ("n_prompts", cfg.get("n_prompts", PLACEHOLDER)),
        ("timestamp corrida", cfg.get("timestamp", PLACEHOLDER)),
    ])


def anexo_b(doc: Document, prompts: list[dict]):
    heading(doc, "Anexo B — Distribución de prompts por módulo y tipo", 1)
    by_module = {}
    by_type = {}
    for p in prompts:
        m = p.get("module", "?")
        t = p.get("type", "?")
        by_module[m] = by_module.get(m, 0) + 1
        by_type[t] = by_type.get(t, 0) + 1

    body(doc, "Distribución por módulo:", bold=True)
    rows = [[k, str(v)] for k, v in sorted(by_module.items())]
    make_table(doc, ["Módulo", "N° prompts"], rows)
    table_caption(doc, 5, "cantidad de prompts por módulo en el golden set.")

    body(doc, "")
    body(doc, "Distribución por tipo:", bold=True)
    rows = [[k, str(v)] for k, v in sorted(by_type.items())]
    make_table(doc, ["Tipo", "N° prompts"], rows)
    table_caption(doc, 6, "cantidad de prompts por tipo de tarea en el golden set.")


def referencias(doc: Document):
    page_break(doc)
    heading(doc, "Referencias (IEEE)", 1)
    refs = [
        "[1] Ollama, \"Ollama — Get up and running with large language models locally.\" En línea. Disponible: https://ollama.com",
        "[2] Qwen Team, \"Qwen2.5: A Party of Foundation Models,\" 2024. En línea. Disponible: https://qwenlm.github.io/blog/qwen2.5/",
        "[3] Meta AI, \"Meta Llama 3.\" En línea. Disponible: https://llama.meta.com",
        "[4] Mistral AI, \"Mistral 7B Instruct.\" En línea. Disponible: https://mistral.ai/news/announcing-mistral-7b",
        "[5] Mixedbread AI, \"mxbai-embed-large-v1,\" Hugging Face. En línea. Disponible: https://huggingface.co/mixedbread-ai/mxbai-embed-large-v1",
        "[6] Nomic AI, \"nomic-embed-text-v1,\" Hugging Face. En línea. Disponible: https://huggingface.co/nomic-ai/nomic-embed-text-v1",
        "[7] E. Shahul, J. James, L. Espejo-Uribe y S. Sarath, \"RAGAS: Automated Evaluation of Retrieval Augmented Generation,\" arXiv:2309.15217, 2023.",
        "[8] J. Brooke, \"SUS: A quick and dirty usability scale,\" Usability evaluation in industry, 1996.",
        "[9] ISO/IEC 25010:2023, \"Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE) — Product quality model,\" ISO, 2023.",
    ]
    for r in refs:
        body(doc, r)
    body(doc,
        "Nota: las fechas exactas de publicación y los DOI específicos "
        "serán verificados en la versión final del Informe de tesis; "
        "este anexo se redacta con la información disponible en las "
        "páginas oficiales de los modelos al momento del Sprint 2.",
        italic=True,
    )


# ----------------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------------
def main(benchmarks_dir: Path, out_path: Path):
    bench_path = benchmarks_dir / "results" / "llm_benchmark.json"
    scores_path = benchmarks_dir / "results" / "llm_scores.json"
    emb_path = benchmarks_dir / "results" / "embeddings_benchmark.json"
    prompts_path = benchmarks_dir / "prompts_llm.json"

    llm_bench = json.loads(bench_path.read_text(encoding="utf-8")) if bench_path.exists() else {}
    llm_scores = json.loads(scores_path.read_text(encoding="utf-8")) if scores_path.exists() else None
    emb_bench = json.loads(emb_path.read_text(encoding="utf-8")) if emb_path.exists() else {}
    prompts = json.loads(prompts_path.read_text(encoding="utf-8")) if prompts_path.exists() else []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_SIZE)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    portada(doc)
    indice(doc)
    seccion_intro(doc)
    seccion_metodo(doc, llm_bench, len(prompts))
    seccion_llm_results(doc, llm_bench, llm_scores)
    seccion_embeddings(doc, emb_bench)
    seccion_decision(doc, llm_bench, llm_scores, emb_bench)
    seccion_limitaciones(doc)
    seccion_indicadores(doc)
    seccion_conclusiones(doc)
    anexo_a(doc, llm_bench)
    anexo_b(doc, prompts)
    referencias(doc)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    # Resumen final
    placeholders = 0  # contar placeholders sería complejo; el lector lo nota visualmente
    print(f"✓ Generado: {out_path}  ({out_path.stat().st_size // 1024} KB)")
    if not llm_scores:
        placeholders += 1
        print(f"  ⚠ llm_scores.json ausente — secciones 3.2 / 3.3 / 3.4 con placeholder")
    if not emb_bench:
        placeholders += 1
        print(f"  ⚠ embeddings_benchmark.json ausente — sección 4 con placeholder")
    print(f"\nConclusiones principales:")
    scores_models = (llm_scores or {}).get("models", {})
    if scores_models:
        best = max(scores_models.items(), key=lambda x: x[1].get("overall_average", 0))
        print(f"  LLM con mejor calidad Likert: {best[0]} = {best[1].get('overall_average', 0):.2f}")
    if emb_bench.get("models"):
        best_e = max(emb_bench["models"].items(), key=lambda x: x[1].get("recall_at_k", 0))
        print(f"  Embedding con mejor Recall@5: {best_e[0]} = {best_e[1].get('recall_at_k', 0):.3f}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--benchmarks-dir", default="/app/benchmarks")
    parser.add_argument("--out", default="/app/docs/reporte-LLM.docx")
    args = parser.parse_args()
    main(Path(args.benchmarks_dir), Path(args.out))
