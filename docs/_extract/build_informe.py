# -*- coding: utf-8 -*-
# Build the extended (example-style) thesis pre-informe as a formatted .docx.
# ponytail: one self-contained build script; python-docx already a project dep.
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = Path(r"C:\tutor-ia-rfa\docs")
MEDIA = DOCS / "_extract" / "media"
OUT = DOCS / "Informe-Final-Tesis-Zavaleta.docx"

doc = Document()

# ---------- global styles ----------
def set_font(style, name="Times New Roman", size=11, bold=False, color=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(a), name)

normal = doc.styles['Normal']
set_font(normal, size=11)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.space_after = Pt(6)

BLACK = RGBColor(0, 0, 0)
for h, sz in (('Heading 1', 14), ('Heading 2', 13), ('Heading 3', 12), ('Heading 4', 11)):
    set_font(doc.styles[h], size=sz, bold=True, color=BLACK)
    doc.styles[h].paragraph_format.space_before = Pt(12)
    doc.styles[h].paragraph_format.space_after = Pt(6)
    doc.styles[h].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    doc.styles[h].paragraph_format.keep_with_next = True

# margins (thesis: left 3cm for binding)
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)

# footer page number (centered)
def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    for t, txt in (('begin', None), ('instr', 'PAGE'), ('end', None)):
        if t == 'instr':
            e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = ' PAGE '
        else:
            e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), t)
        r._r.append(e)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

add_page_number(doc.sections[0])

# ---------- helpers ----------
TBL = [0]; FIG = [0]

def para(text="", align='justify', bold=False, italic=False, size=11, space_after=6, center=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    if center or align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        r.font.name = 'Times New Roman'
    return p

def heading(text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Times New Roman'; run.font.color.rgb = BLACK
    return h

def bullets(items, style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True; r.font.name = 'Times New Roman'
            r2 = p.add_run(it[1]); r2.font.name = 'Times New Roman'
        else:
            r = p.add_run(it); r.font.name = 'Times New Roman'

def labeled(label, text):
    para(bold=False)  # noop guard removed below
def label_run(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label); r.bold = True; r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.font.name = 'Times New Roman'
    return p

def table(headers, rows, caption, col_w=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.name = 'Times New Roman'
        shade(hdr[i], 'D9E2F3')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.size = Pt(10); r.font.name = 'Times New Roman'
    if col_w:
        for i, w in enumerate(col_w):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    TBL[0] += 1
    cap = para(f"Tabla {TBL[0]}. {caption}", center=True, italic=True, size=9, space_after=10)
    return t

def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcpr.append(shd)

def figure(img, caption, width_in):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run()
    r.add_picture(str(MEDIA / img), width=Inches(width_in))
    FIG[0] += 1
    para(f"Figura {FIG[0]}. {caption}", center=True, italic=True, size=9, space_after=10)

def toc():
    p = doc.add_paragraph()
    r = p.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = 'Haga clic derecho sobre este texto y elija «Actualizar campo» para generar el índice automático.'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    for e in (b, instr, sep, t, end):
        r._r.append(e)

def center_line(text, bold=False, size=12, space=2, caps=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space); p.paragraph_format.line_spacing = 1.15
    rr = p.add_run(text.upper() if caps else text)
    rr.bold = bold; rr.font.size = Pt(size); rr.font.name = 'Times New Roman'
    return p

# ======================================================================
# PORTADA
# ======================================================================
center_line("UNIVERSIDAD CATÓLICA SANTO TORIBIO DE MOGROVEJO", bold=True, size=13, space=2)
center_line("FACULTAD DE INGENIERÍA", bold=True, size=12, space=2)
center_line("ESCUELA DE INGENIERÍA DE SISTEMAS Y COMPUTACIÓN", bold=True, size=12, space=8)
plogo = doc.add_paragraph(); plogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
plogo.add_run().add_picture(str(MEDIA / 'image1.jpeg'), width=Inches(1.5))
center_line("", space=8)
center_line("TUTOR CON IA GENERATIVA EN LA ASIGNATURA APLICACIONES MÓVILES "
            "PARA MEJORAR EL RENDIMIENTO ACADÉMICO DE ESTUDIANTES DEL IESTP “RFA”, CHICLAYO",
            bold=True, size=14, space=14)
center_line("PRE-INFORME DE TESIS", bold=True, size=12, space=2)
center_line("PARA EVIDENCIAR LA EJECUCIÓN DE TESIS", size=11, space=2)
center_line("EN EL CURSO DE SEMINARIO DE TESIS I", size=11, space=16)
center_line("AUTOR", bold=True, size=12, space=2)
center_line("Zavaleta Marcelo, Roger Alessandro", size=11, space=1)
center_line("DNI: 70469567", size=11, space=14)
center_line("ASESORA", bold=True, size=12, space=2)
center_line("Mg. Reyes Burgos, Karla", size=11, space=1)
center_line("https://orcid.org/0000-0002-9650-4427", size=10, space=16)
center_line("Chiclayo, 2026", bold=True, size=12, space=2)
doc.add_page_break()

# ======================================================================
# INFORMACIÓN GENERAL
# ======================================================================
heading("INFORMACIÓN GENERAL", 1)
ig = doc.add_table(rows=0, cols=2); ig.style = 'Table Grid'
info = [
    ("1. Facultad y Escuela", "Facultad de Ingeniería – Escuela de Ingeniería de Sistemas y Computación"),
    ("2. Título del informe de tesis", "Tutor con IA generativa en la asignatura aplicaciones móviles para mejorar el rendimiento académico de estudiantes del IESTP “RFA”, Chiclayo"),
    ("3. Autor y firma", "Zavaleta Marcelo, Roger Alessandro — DNI 70469567"),
    ("4. Asesora y firma", "Mg. Reyes Burgos, Karla — ORCID 0000-0002-9650-4427"),
    ("5. Línea y área de la investigación", "Desarrollo e innovación tecnológica – Inteligencia artificial"),
    ("6. Fecha de presentación", "21 de junio de 2026"),
]
for k, v in info:
    c = ig.add_row().cells
    c[0].text = ''; r = c[0].paragraphs[0].add_run(k); r.bold = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    c[1].text = ''; r2 = c[1].paragraphs[0].add_run(v); r2.font.size = Pt(11); r2.font.name = 'Times New Roman'
    c[0].width = Cm(5.5); c[1].width = Cm(10.5)
doc.add_page_break()

# ======================================================================
# ÍNDICE
# ======================================================================
heading("ÍNDICE", 1)
toc()
doc.add_page_break()

# ======================================================================
# I. RESULTADOS
# ======================================================================
heading("I. RESULTADOS", 1)
para("Los resultados se presentan exponiendo los hallazgos y productos del desarrollo del "
     "tutor con IA generativa para la asignatura Aplicaciones Móviles del Instituto de Educación "
     "Superior Tecnológico Público “República Federal de Alemania” (IESTP “RFA”), de Chiclayo. "
     "El proyecto se estructuró con la metodología ágil SCRUM como marco principal de gestión, "
     "complementada con CRISP-DM exclusivamente para las fases asociadas al componente de "
     "aprendizaje de máquina del pipeline de Recuperación Aumentada por Generación (RAG). El "
     "trabajo se organizó en doce sprints distribuidos en seis iteraciones del modelo CRISP-DM. "
     "La presentación se ordena en dos planos complementarios: el plano metodológico (apartado 1.1), "
     "que recorre las seis iteraciones del cronograma y los sprints que las componen; y el plano de "
     "objetivos (apartado 1.2), que evidencia el cumplimiento de cada uno de los cinco objetivos "
     "específicos frente a sus indicadores objetivamente verificables.")
para("El objetivo general de la investigación fue desarrollar un tutor con IA generativa aplicado a "
     "la asignatura de Aplicaciones Móviles para mejorar el rendimiento académico de los estudiantes "
     "del IESTP “RFA”. Este objetivo se operacionalizó en cinco objetivos específicos: (OE1) seleccionar "
     "el modelo de lenguaje grande (LLM) y el modelo de embeddings para la generación y recuperación de "
     "respuestas en español del dominio; (OE2) validar la precisión de recuperación y la fidelidad de "
     "generación del pipeline RAG mediante el framework RAGAS; (OE3) desplegar el sistema sobre Google "
     "Compute Engine en contenedores Docker, garantizando rendimiento, disponibilidad y trazabilidad; "
     "(OE4) contrastar la mejora del rendimiento académico mediante un diseño pre-experimental con "
     "pretest y postest; y (OE5) validar la adecuación funcional del sistema conforme a la norma "
     "ISO/IEC 25010:2023.")

# ---------------------------------------------------------------- 1.1
heading("1.1. En base a la metodología utilizada", 2)
para("El sistema se desarrolló empleando SCRUM para la gestión general del proyecto y CRISP-DM como "
     "marco que estructura el ciclo de vida del componente de aprendizaje de máquina del pipeline RAG. "
     "Los doce sprints se distribuyeron en seis iteraciones CRISP-DM: Iteración #1 Comprensión del negocio "
     "(Sprint 1); Iteración #2 Comprensión de los datos (Sprint 2); Iteración #3 Preparación de los datos y "
     "modelado (Sprints 3–6); Iteración #4 Evaluación (Sprint 7); Iteración #5 Despliegue (Sprints 8–10); e "
     "Iteración #6 Validación con usuarios y pilotaje (Sprints 11–12). SCRUM permitió entregas funcionales "
     "incrementales y la revisión continua; CRISP-DM aportó el marco que justifica las fases del pipeline RAG "
     "y la lógica de validación de los modelos seleccionados (RAGAS) y del sistema integrado (ISO/IEC 25010:2023 "
     "y diseño pre-experimental con pretest/postest).")

# ================= 1.1.1 Iteración 1 — Comprensión del negocio (EXPANDED) =================
heading("1.1.1. Iteración #1: Comprensión del negocio (Sprint 1)", 3)
para("La Iteración #1 (Sprint 1, 23 – 29 de marzo de 2026) estableció la base documental y técnica del "
     "proyecto. Siguiendo la fase de Comprensión del negocio de CRISP-DM, se determinaron los objetivos del "
     "negocio, se evaluó la situación actual de la asignatura, se establecieron los objetivos técnicos (de "
     "minería de datos) y se produjo el plan del proyecto. Como insumos de diseño se formalizó la "
     "Especificación de Requisitos del Software (ERS), se diseñó la arquitectura de microservicios "
     "contenedorizados y se construyeron los dos primeros modelos del Sistema de Tutoría Inteligente (dominio "
     "y pedagógico).")

heading("1.1.1.1. Determinar los objetivos del negocio", 4)
label_run("Objetivos principales: ", "")
bullets([
    "Mejorar el rendimiento académico de los estudiantes en la asignatura Aplicaciones Móviles (Android/Kotlin) del IESTP “RFA”.",
    "Brindar tutoría personalizada disponible de forma continua mediante IA generativa con conocimiento del dominio, recuperado del sílabo oficial (RAG).",
    "Reducir la brecha de acompañamiento docente fuera del aula, ofreciendo retroalimentación inmediata y trazable a las fuentes del curso.",
    "Operar con una infraestructura 100 % privada (LLM autohospedado), sin recurrir a APIs pagas y preservando la privacidad de los datos de los estudiantes.",
])
label_run("Criterios de éxito: ", "")
bullets([
    "Mejora estadísticamente significativa del rendimiento académico (pretest/postest, p < 0.05).",
    "Pipeline RAG validado con métricas de recuperación y generación por encima de umbral (RAGAS).",
    "Sistema desplegado, accesible y con trazabilidad de las respuestas al corpus instruccional.",
    "Adecuación funcional del sistema conforme a ISO/IEC 25010:2023 por encima de los umbrales definidos.",
])

heading("1.1.1.2. Evaluar la situación actual", 4)
label_run("Análisis de la realidad estudiada. ",
     "El IESTP “República Federal de Alemania” (Chiclayo) imparte la asignatura Aplicaciones Móviles dentro de "
     "su programa de Tecnologías de la Información. La enseñanza de la programación —en particular Kotlin y el "
     "desarrollo Android— presenta una dificultad reconocida en la literatura, con altas tasas de reprobación "
     "en cursos introductorios de programación [1], [2]. En el contexto nacional, los resultados del Perú en "
     "PISA 2022 evidencian brechas persistentes en competencias fundamentales [15], y la enseñanza técnica "
     "superior enfrenta limitaciones de acompañamiento personalizado [14]. Los Sistemas de Tutoría Inteligente "
     "han mostrado eficacia comparable a la tutoría humana individual [3], [4], y la IA generativa con RAG "
     "permite acotar las respuestas al material del curso, reduciendo las alucinaciones [6].")
label_run("Deficiencias identificadas: ", "")
bullets([
    "Acompañamiento docente limitado fuera del horario de clase y una relación docente–estudiante que dificulta la atención individualizada.",
    "Ausencia de retroalimentación inmediata sobre los ejercicios de programación del estudiante.",
    "Recursos de estudio genéricos, no siempre alineados al sílabo vigente de la asignatura.",
    "Dependencia del docente para resolver dudas puntuales, con la consiguiente demora en la realimentación.",
])
label_run("Necesidades para el producto acreditable: ", "")
bullets([
    "Un tutor conversacional con conocimiento del dominio, capaz de responder consultas citando el material del curso.",
    "Evaluación adaptativa que ubique al estudiante en un nivel y ajuste la dificultad del contenido.",
    "Retroalimentación automática del código del estudiante con criterios pedagógicos.",
    "Trazabilidad de las respuestas (citas verificables) y operación privada del modelo de lenguaje.",
])
label_run("Funciones del producto acreditable: ", "")
bullets([
    "Tutor conversacional basado en RAG, con respuestas en español peruano y citas a las fuentes del corpus.",
    "Cuestionarios (quizzes) generados por IA con retroalimentación y banco de respaldo.",
    "Desafíos de programación con evaluación automática del código y feedback adaptado al nivel.",
    "Evaluación de entrada y nivelación dinámica del estudiante (principiante, intermedio, avanzado).",
    "Seguimiento del progreso gamificado y panel de administración del corpus (RAG).",
])
label_run("Recursos disponibles: ", "")
bullets([
    "Personal: tesista (desarrollador), asesora de la USAT (Mg. Reyes Burgos, Karla) y coordinador del piloto en el IESTP “RFA” (Téc. Xavier Benites Marín).",
    "Hardware: equipo de desarrollo con GPU (NVIDIA RTX 4070) y una máquina virtual e2-standard-4 (16 GB RAM) en Google Compute Engine para el despliegue.",
    "Software: pila íntegramente de código abierto (FastAPI, React, Ollama, PostgreSQL + pgvector, Redis, Docker, Caddy).",
    "Datos: corpus instruccional derivado del sílabo oficial 2026-I de la asignatura Aplicaciones Móviles.",
])
para("La especificación detallada de requisitos consolidó 52 requisitos (33 funcionales priorizados y 19 no "
     "funcionales) distribuidos en ocho módulos del sistema y alineados con ISO/IEC 25010:2023. La Tabla 1 "
     "sintetiza la distribución por categoría y prioridad, y la Tabla 2 traza su correspondencia con los cinco "
     "objetivos específicos vigentes (matriz V.2.1).")
table(["Categoría", "Total", "Prioridad alta", "Prioridad media", "Prioridad baja"],
      [["Requisitos Funcionales (RF)", "33", "22", "10", "1"],
       ["Requisitos No Funcionales (RNF)", "19", "13", "6", "0"],
       ["TOTAL", "52", "35", "16", "1"]],
      "Síntesis de requisitos especificados en la ERS v1.0.")
table(["Objetivo específico (V.2.1)", "RF asociados", "RNF asociados"],
      [["OE1: Selección de LLM y de embeddings", "Estudio comparativo externo al software", "RNF-02.3, RNF-03.6, RNF-07.2"],
       ["OE2: Validación RAGAS del pipeline RAG", "RF-19, RF-20, RF-21, RF-22", "RNF-02.3, RNF-02.5"],
       ["OE3: Despliegue en GCE con trazabilidad", "RF-07 a RF-29, RF-33", "RNF-01, RNF-02, RNF-05, RNF-06"],
       ["OE4: Mejora del rendimiento (pretest/postest)", "Instrumento externo al software", "—"],
       ["OE5: Adecuación funcional ISO/IEC 25010", "RF-18, RF-25, RF-27, RF-29, RF-32", "RNF-01, RNF-02, RNF-03, RNF-04"]],
      "Trazabilidad de requisitos a los cinco objetivos específicos vigentes (matriz V.2.1).")
label_run("Presupuesto del proyecto. ",
     "El proyecto se ejecutó con un presupuesto estimado de S/ 3 170.00, concentrado en la infraestructura de "
     "cómputo en la nube y la conectividad necesaria para el despliegue y el pilotaje. La Tabla 3 detalla la "
     "asignación estimada por rubro.")
table(["Rubro", "Descripción", "Monto (S/)"],
      [["Infraestructura de cómputo", "Máquina virtual e2-standard-4 en Google Compute Engine (despliegue y pruebas)", "1 600.00"],
       ["Cómputo acelerado (GPU)", "Instancia con GPU para la medición de rendimiento del pipeline (OE3)", "600.00"],
       ["Dominio y certificados", "Registro del dominio y TLS (Let’s Encrypt, sin costo de certificado)", "120.00"],
       ["Conectividad e internet", "Servicio de internet durante el desarrollo y el piloto", "450.00"],
       ["Materiales y contingencia", "Materiales del piloto, impresiones y reserva de contingencia", "400.00"],
       ["TOTAL", "", "3 170.00"]],
      "Presupuesto estimado del proyecto por rubro.")
label_run("Supuestos y restricciones: ", "")
bullets([
    "El modelo de lenguaje opera 100 % de forma privada mediante Ollama autohospedado; no se emplean APIs pagas (OpenAI, Anthropic, Gemini).",
    "El conocimiento del dominio se incorpora exclusivamente mediante RAG sobre el corpus del sílabo; no hay fine-tuning del modelo.",
    "La pila tecnológica es cerrada: una única VM en Google Compute Engine y Firebase Hosting para el frontend.",
    "La interfaz de usuario se redacta en español peruano; el código fuente se escribe en inglés.",
])
label_run("Riesgos y contingencias: ", "")
bullets([
    "Riesgo: caída del motor Ollama. Contingencia: degradación elegante con bancos de respaldo (quiz desde banco docente, coding desde catálogo, evaluación de entrada desde banco).",
    "Riesgo: la VM del piloto es CPU-only y no alcanza el rendimiento de inferencia objetivo. Contingencia: medición del rendimiento sobre una instancia con GPU, manteniendo el mismo backend de producción.",
    "Riesgo: muestra reducida del piloto (10–15 estudiantes). Contingencia: aplicación censal a la cohorte 2026-I completa (n = 49).",
])

heading("1.1.1.3. Establecer los objetivos de minería de datos", 4)
label_run("Objetivos técnicos: ", "")
bullets([
    "Construir un pipeline RAG sobre el corpus del sílabo capaz de recuperar contexto pertinente y generar respuestas fieles, con trazabilidad explícita de las fuentes.",
    "Seleccionar el LLM open-source y el modelo de embeddings con mejor desempeño en español sobre el dominio de Aplicaciones Móviles.",
    "Garantizar que la recuperación y la generación superen umbrales de calidad medibles (RAGAS) sin recurrir a fine-tuning.",
])
label_run("Criterios de éxito técnicos: ", "")
bullets([
    "Recuperación: Context Precision ≥ 0.70 y Context Recall ≥ 0.75.",
    "Generación (clase de modelo 7B autohospedado): Faithfulness ≥ 0.65, Answer Relevancy ≥ 0.65 y Answer Correctness ≥ 0.55.",
    "Selección del LLM (Accuracy ≥ 0.70, Likert ≥ 4.0) y de embeddings (nDCG@10 ≥ 0.55, Recall@5 ≥ 0.70, MRR@10 ≥ 0.65).",
])

heading("1.1.1.4. Producir el plan del proyecto", 4)
para("El plan del proyecto se organizó en doce sprints de SCRUM agrupados en seis iteraciones CRISP-DM, con "
     "una duración total de 106 días (23 de marzo – 10 de julio de 2026). La Tabla 4 resume el cronograma y su "
     "correspondencia con los objetivos específicos.")
table(["Iteración CRISP-DM", "Sprints", "Periodo", "Foco / objetivo"],
      [["#1 Comprensión del negocio", "S1", "23–29 mar", "ERS, arquitectura C4 y modelos de diseño (insumo OE3)"],
       ["#2 Comprensión de los datos", "S2", "30 mar–05 abr", "Selección de LLM y de embeddings (OE1)"],
       ["#3 Preparación y modelado", "S3–S6", "06–24 abr", "Pipeline RAG, backend, frontend, RAGAS preliminar (OE2/OE3)"],
       ["#4 Evaluación", "S7", "25–29 abr", "Reranking y validación formal RAGAS (OE2)"],
       ["#5 Despliegue", "S8–S10", "30 abr–20 may", "Despliegue en GCE, contenido e instrumento (OE3)"],
       ["#6 Validación y pilotaje", "S11–S12", "21 may–10 jul", "ISO/IEC 25010 (OE5) y pretest/postest (OE4)"]],
      "Cronograma del proyecto: doce sprints SCRUM en seis iteraciones CRISP-DM.")
label_run("Herramientas: ",
     "Docker 24.0 + Docker Compose 2.20 (orquestación), PostgreSQL 16 + pgvector (base relacional y vectorial), "
     "SQLAlchemy 2.x async + asyncpg (ORM y driver asíncrono), Alembic (versionado del esquema) y Git (control de versiones).")
label_run("Artefactos / entregables: ", "")
bullets([
    "Especificación de Requisitos del Software con 52 requisitos.",
    "Archivos docker-compose.yml y .env.example con la orquestación inicial del sistema.",
    "Modelos ORM iniciales y migración inicial de Alembic; arquitectura C4 y modelos de dominio y pedagógico del STI.",
])

# ================= 1.1.2 Iteración 2 — Comprensión de los datos (EXPANDED) =================
heading("1.1.2. Iteración #2: Comprensión de los datos (Sprint 2)", 3)
para("La Iteración #2 (Sprint 2, 30 de marzo – 05 de abril de 2026) abordó la fase de Comprensión de los datos "
     "de CRISP-DM: la recopilación, descripción, exploración y verificación de la calidad del corpus "
     "instruccional, y la evaluación comparativa de los modelos de lenguaje y de embeddings que alimentarán el "
     "pipeline RAG. Con ello se cerró la selección técnica del OE1 (Resultados R1.1 y R1.2) y se completaron los "
     "dos modelos restantes del STI (modelo del estudiante y modelo de interacción).")

heading("1.1.2.1. Recopilación, descripción y exploración del corpus", 4)
label_run("a) Recopilación de datos. ",
     "El corpus del sistema se construyó a partir del sílabo oficial 2026-I de la asignatura Aplicaciones "
     "Móviles del IESTP “RFA”. Se reunieron los documentos instruccionales de los cinco módulos del curso en "
     "formatos PDF, DOCX y Markdown, que constituyen la única fuente de conocimiento del dominio del tutor "
     "(no se emplean fuentes externas: el modelo solo responde desde este corpus).")
label_run("b) Descripción de los datos. ",
     "El corpus se organiza jerárquicamente en cinco módulos y veintidós temas, según se resume en la Tabla 5. "
     "Tras el preprocesamiento, el contenido se segmenta en fragmentos (chunks) y se vectoriza con embeddings de "
     "1 024 dimensiones; en el entorno de producción se indexaron 3 388 fragmentos en PostgreSQL con la extensión "
     "pgvector.")
table(["Módulo", "Contenido principal", "N° de temas"],
      [["M1", "Fundamentos del desarrollo de aplicaciones móviles", "4"],
       ["M2", "Lenguaje de programación Kotlin", "5"],
       ["M3", "Diseño de interfaces de usuario (UI)", "4"],
       ["M4", "Componentes de la aplicación y manejo de datos", "5"],
       ["M5", "Temas avanzados y despliegue", "4"],
       ["TOTAL", "5 módulos", "22 temas"]],
      "Estructura del corpus instruccional derivado del sílabo de Aplicaciones Móviles (2026-I).")
label_run("c) Exploración inicial de los datos. ",
     "Se analizó la distribución del contenido por módulo y la cobertura temática del sílabo. La exploración "
     "confirmó que el corpus abarca desde los fundamentos del desarrollo móvil (M1) hasta los temas avanzados de "
     "despliegue (M5), con una densidad de contenido suficiente para sostener consultas conceptuales, de código y "
     "de aplicación en cada módulo.")
label_run("d) Verificación de la calidad de los datos. ",
     "Se verificó la integridad y la consistencia de los documentos, normalizando los espacios en blanco, "
     "limitando los saltos de línea consecutivos y eliminando fragmentos vacíos o duplicados. Se confirmó que la "
     "segmentación preserva la jerarquía del sílabo (identificadores de módulo, tema y subtema en los metadatos), "
     "condición necesaria para la trazabilidad posterior de las citas.")

heading("1.1.2.2. Selección comparativa del LLM y del modelo de embeddings", 4)
para("Sobre el corpus descrito se ejecutó la evaluación comparativa de tres modelos LLM open-source servidos por "
     "Ollama (qwen2.5:7b-instruct-q4_K_M, llama3:8b-instruct-q4_K_M y mistral:7b-instruct-q4_K_M) sobre un golden "
     "set de 50 prompts en español, midiendo calidad con una rúbrica Likert 1–5, consumo de memoria y latencia "
     "(Tablas 6 y 7). En paralelo se compararon dos modelos de embeddings (mxbai-embed-large y nomic-embed-text) "
     "para similitud semántica en español sobre 20 consultas representativas del corpus (Tabla 8).")
table(["Modelo", "Tamaño (GB)", "Lat. avg (s)", "Lat. p95 (s)", "Tokens/s", "Calidad (1-5)", "VRAM (GB)"],
      [["qwen2.5:7b", "4.59", "5.91", "8.66", "107.0", "4.85", "4.59"],
       ["llama3:8b", "5.09", "4.71", "7.40", "109.6", "4.84", "5.09"],
       ["mistral:7b", "4.79", "3.90", "6.79", "120.2", "4.28", "4.79"]],
      "Comparativa global de rendimiento de los tres LLM evaluados (50 prompts c/u).")
table(["Modelo", "Exactitud", "Fluidez", "Sin alucinaciones", "Pedagogía", "Promedio"],
      [["qwen2.5:7b", "4.72", "4.86", "5.00", "4.82", "4.85"],
       ["llama3:8b", "4.72", "4.88", "4.92", "4.86", "4.84"],
       ["mistral:7b", "3.90", "4.56", "4.60", "4.04", "4.28"]],
      "Promedios Likert por criterio de la rúbrica (juez LLM).")
table(["Modelo", "Dimensiones", "Recall@5", "MRR", "Latencia/query (ms)"],
      [["mxbai-embed-large", "1024", "0.550", "0.453", "173.4"],
       ["nomic-embed-text", "768", "0.350", "0.305", "131.0"]],
      "Comparativa de modelos de embeddings sobre el corpus M1–M3 (20 queries / 163 chunks).")
para("La selección final recayó en qwen2.5:7b-instruct-q4_K_M como LLM, por su valoración de calidad superior "
     "(4.85) y su ausencia total de alucinaciones (5.00), y en mxbai-embed-large como modelo de embeddings, por su "
     "precisión de recuperación marcadamente superior (Recall@5 0.550 frente a 0.350). La validación de estos "
     "modelos contra los umbrales del OE1 se realiza sobre el pipeline de producción y se reporta en el apartado 1.2.1.")
label_run("Herramientas: ",
     "Figma (mockups UI/UX), PlantUML (diagramas de secuencia UML), Draw.io (diagrama ER y grafo del modelo de "
     "dominio), Ollama (servidor local de modelos), qwen2.5:7b-instruct-q4_K_M (LLM) y mxbai-embed-large "
     "(embeddings de 1 024 dimensiones).")
label_run("Artefactos / entregables: ", "")
bullets([
    "Reporte comparativo de modelos LLM y de embeddings.",
    "Modelfile personalizado del LLM (ollama/modelfile-qwen2.5) con el system prompt pedagógico del tutor.",
    "Modelos del estudiante y de interacción del STI especificados.",
])

# ================= 1.1.3 Iteración 3 — Preparación de los datos y modelado =================
heading("1.1.3. Iteración #3: Preparación de los datos y modelado (Sprints 3–6)", 3)
para("La Iteración #3 (Sprints 3–6, 06 – 24 de abril de 2026) construyó el núcleo funcional del sistema: el "
     "pipeline RAG operativo, el backend REST, el frontend integrado de extremo a extremo y la primera "
     "validación diagnóstica del pipeline con RAGAS. Cubre las fases CRISP-DM de Preparación de los datos y "
     "Modelado del componente de IA.")

heading("1.1.3.1. Sprint 3: Pipeline RAG, ingesta, chunking e indexación vectorial", 4)
para("El Sprint 3 (06 – 12 de abril de 2026) construyó el pipeline RAG operativo sobre el corpus de los Módulos "
     "1–3 del sílabo: un módulo de chunking con RecursiveCharacterTextSplitter (fragmentos de 500 tokens, "
     "solapamiento de 50) que preserva la jerarquía del sílabo en los metadatos; el cálculo de embeddings con "
     "mxbai-embed-large y su persistencia en PostgreSQL con pgvector (índice IVFFlat, similitud coseno); un "
     "orquestador que recupera el contexto (top-k = 5), arma el prompt aumentado con historial y genera la "
     "respuesta vía la API de Ollama con trazabilidad de fuentes; y un servicio de ingesta (IngestService) como "
     "BackgroundTask para parsear PDF, DOCX y Markdown. Los parámetros del pipeline al cierre del sprint se "
     "resumen en la Tabla 9 y el flujo de una consulta se ilustra en la Figura 1.")
table(["Parámetro", "Implementación inicial", "Comentario"],
      [["Modelo LLM", "qwen2.5:7b-instruct-q4_K_M", "Confirmado en Sprint 2."],
       ["Modelo de embeddings", "mxbai-embed-large (1 024 dim)", "Confirmado en Sprint 2."],
       ["Estrategia de chunking", "RecursiveCharacterTextSplitter (500 / 50)", "Iteración baseline."],
       ["Umbral de similitud (coseno)", "0.70", "Iteración baseline; se ajusta en Sprint 6."],
       ["Top-k del retrieval", "5", "Iteración baseline."],
       ["num_ctx del LLM", "4 096", "Iteración baseline."],
       ["Temperatura del LLM", "0.3", "Iteración baseline."],
       ["TTL de la caché RAG", "3 600 s (1 h)", "Confirmado."]],
      "Parámetros del pipeline RAG al cierre del Sprint 3.")
figure('image2.png', "Flujo del pipeline RAG para una consulta del estudiante.", 3.1)
label_run("Artefactos / entregables: ", "")
bullets([
    "Módulos del pipeline RAG: ingestor, retriever, generator y orchestrator.",
    "Servicios de dominio: IngestService, LLMService y RAGService.",
    "Corpus de los Módulos 1–3 indexado en PostgreSQL + pgvector (tabla document_chunks con índice IVFFlat).",
])

heading("1.1.3.2. Sprint 4: Backend FastAPI con autenticación JWT y endpoints REST", 4)
para("El Sprint 4 (13 – 17 de abril de 2026) implementó el backend completo: autenticación JWT (access token de "
     "60 min y refresh token de 7 días, contraseñas con bcrypt), control de acceso por roles (estudiante y "
     "administrador), diez routers REST con más de veinticinco endpoints bajo el prefijo /api/v1, y middlewares "
     "transversales (rate limiting con slowapi de 20 consultas/hora en el chat y logging estructurado en JSON con "
     "loguru). La Tabla 10 resume los routers y su función.")
table(["Router", "Endpoints principales", "Función"],
      [["/auth", "POST /register · /login · /refresh · /logout", "Registro, autenticación y ciclo de vida de tokens."],
       ["/users", "GET /me · PUT /me · GET /me/level", "Perfil y nivel asignado por la evaluación de entrada."],
       ["/modules", "GET / · GET /{id}", "Listado y detalle de módulos con estado de desbloqueo."],
       ["/topics", "GET /{id} · POST /{id}/complete", "Contenido, video y marcado de completitud."],
       ["/quiz", "GET /{topic_id} · POST /submit", "Preguntas generadas por IA y cálculo de puntaje."],
       ["/chat", "POST /sessions · /messages · GET /remaining", "Tutor conversacional con RAG y rate limiting."],
       ["/dashboard", "GET /", "Nivel, progreso global, logros y recomendaciones."],
       ["/coding", "GET /topic/{id} · POST /{id}/submit", "Desafíos de código adaptados y feedback de la IA."],
       ["/assessment", "POST /start · /submit", "Evaluación de entrada que clasifica en un nivel."],
       ["/admin", "POST /corpus/upload · GET /users", "Gestión del corpus y supervisión de estudiantes."]],
      "Routers REST del backend y su función.")
label_run("Artefactos / entregables: ", "")
bullets([
    "Backend FastAPI con 10 routers y más de 25 endpoints REST operativos.",
    "Servicios de dominio: AuthService, ProgressService, LevelingService y ChallengeGenerator.",
    "Módulos de seguridad: JWT, rate limiting (slowapi) y logging (loguru).",
])

heading("1.1.3.3. Sprint 5: Frontend React SPA e integración end-to-end", 4)
para("El Sprint 5 (18 – 22 de abril de 2026) desarrolló el frontend de página única (SPA) con React 18, "
     "TypeScript estricto y shadcn/ui, integrándolo con el backend REST. Al cierre del sprint el sistema funcionó "
     "de extremo a extremo en entorno de desarrollo: el flujo estudiante → consulta → respuesta con citas operó "
     "correctamente. La vista de contenedores del sistema (modelo C4, nivel 2) se inventaría en la Tabla 11 y se "
     "ilustra en la Figura 2.")
table(["Contenedor", "Tecnología", "Responsabilidad principal", "Puerto"],
      [["SPA Frontend", "React 18 + Vite + TS + shadcn/ui", "Interfaz web; sesión JWT y consumo de la API.", "443"],
       ["Reverse Proxy", "Caddy + Let’s Encrypt", "Termina TLS, rutea al backend, renueva certificados.", "443→80"],
       ["Backend API", "FastAPI + Python 3.11", "API REST, JWT, orquestación RAG y BackgroundTasks.", "8000"],
       ["Motor LLM", "Ollama + qwen2.5:7b", "Genera respuestas y embeddings. No expuesto.", "11434"],
       ["Base de datos", "PostgreSQL 16 + pgvector", "Persistencia relacional y vectorial (IVFFlat).", "5432"],
       ["Caché", "Redis 7", "Caché de respuestas RAG (TTL 3 600 s) y rate limit.", "6379"],
       ["Corpus y backups", "Volumen /data en la VM", "Documentos fuente y respaldos de PostgreSQL.", "—"]],
      "Inventario de contenedores del sistema (C4 · Nivel 2).")
figure('image3.png', "Diagrama de contenedores del sistema (C4 · Nivel 2).", 5.6)
label_run("Artefactos / entregables: ", "")
bullets([
    "Frontend React con páginas operativas: Login, Dashboard, Modules, Lesson y Progress.",
    "Componentes de chat: ChatInterface, MessageBubble y SourceCitation (fuentes y % de similitud).",
    "Cliente HTTP centralizado con interceptor de tokens; sistema funcional end-to-end verificado por pruebas manuales.",
])

heading("1.1.3.4. Sprint 6: Validación preliminar del pipeline RAG con RAGAS", 4)
para("El Sprint 6 (23 – 24 de abril de 2026) ejecutó la validación preliminar del pipeline con RAGAS sobre un "
     "golden set inicial de 30 preguntas con ground-truth. Esta iteración diagnóstica fue determinante: reveló "
     "una precisión de contexto (context_precision) baja, del orden de 0.29, con recuperación por similitud "
     "coseno simple, lo que motivó la incorporación de una etapa de reranking con cross-encoder en el pipeline "
     "(Tabla 12).")
table(["Métrica", "Baseline", "v2", "Δ"],
      [["faithfulness", "0.663", "0.716", "+0.053"],
       ["answer_relevancy", "0.863", "0.856", "−0.007"],
       ["context_precision", "0.297", "0.290", "−0.007"],
       ["context_recall", "0.547", "0.619", "+0.072"]],
      "Resultados RAGAS preliminares (30 preguntas).")
para("Hallazgo metodológico: la recuperación por similitud coseno simple no alcanza una precisión de contexto "
     "aceptable para el dominio. La contribución técnica del proyecto —la incorporación de reranking con "
     "cross-encoder y la ampliación del golden set a 50 ítems— se mide en el Sprint 7 (Iteración #4) y se reporta "
     "en el apartado 1.2.2, donde las cinco métricas superan sus umbrales.")

# ================= 1.1.4 Iteración 4 — Evaluación =================
heading("1.1.4. Iteración #4: Evaluación (Sprint 7)", 3)
para("La Iteración #4 (Sprint 7, 25 – 29 de abril de 2026) materializó la contribución técnica del proyecto: se "
     "incorporó una etapa de reranking con cross-encoder sobre los fragmentos recuperados por similitud coseno, y "
     "se ejecutó la validación formal del pipeline RAG con la librería oficial de RAGAS (ragas==0.2.6) sobre un "
     "golden set ampliado a 50 preguntas, empleando un juez LLM independiente (llama3.1:8b) distinto del modelo "
     "generador para evitar el sesgo de autoevaluación. El diseño fue de tipo antes/después controlado: se "
     "midieron las mismas métricas con recuperación coseno simple (línea base) y con coseno + reranking, aislando "
     "el efecto de la mejora (Tabla 13).")
table(["Métrica de recuperación", "Sin rerank (coseno)", "Con rerank (cross-encoder)", "Δ"],
      [["Recall@5", "0.620", "0.720", "+0.100"],
       ["MRR@10", "0.535", "0.684", "+0.149"],
       ["nDCG@10", "0.568", "0.686", "+0.118"]],
      "Mejora de la recuperación atribuible al reranking con cross-encoder (antes/después controlado, 50 ítems).")
para("La mejora consistente en las tres métricas de ordenamiento (Recall@5 +0.100, MRR@10 +0.149, nDCG@10 "
     "+0.118) constituye la evidencia cuantitativa de la contribución técnica: el reranking eleva la calidad de "
     "la recuperación por encima de la línea base de similitud coseno simple. Este incremento es el que permite "
     "que la context_precision pase de ≈ 0.29 (Sprint 6) a 0.876 en la validación formal (apartado 1.2.2).")

# ================= 1.1.5 Iteración 5 — Despliegue =================
heading("1.1.5. Iteración #5: Despliegue (Sprints 8–10)", 3)
para("La Iteración #5 (Sprints 8–10, 30 de abril – 20 de mayo de 2026) consolidó los servicios de soporte, "
     "desplegó el sistema en producción sobre Google Compute Engine y produjo los insumos pedagógicos del estudio "
     "de campo.")
heading("1.1.5.1. Sprint 8: Tareas en segundo plano, caché y provisión de infraestructura", 4)
bullets([
    "Integración de Redis 7 como caché de respuestas del pipeline RAG (TTL de 3 600 s) y como backend del control de tasa.",
    "Tareas programadas con APScheduler para mantenimiento del índice y respaldos periódicos de PostgreSQL.",
    "Provisión de la máquina virtual e2-standard-4 (4 vCPU, 16 GB RAM) en la región us-central1-a de Google Compute Engine.",
    "Endurecimiento básico de la instancia: reglas de firewall, usuario no root y variables de entorno gestionadas mediante archivo .env.",
])
heading("1.1.5.2. Sprint 9: Despliegue productivo en Google Compute Engine", 4)
para("El Sprint 9 (7 – 13 de mayo de 2026) ejecutó el despliegue productivo del tutor en Google Compute Engine, "
     "en contenedores Docker orquestados con Docker Compose, con terminación TLS gestionada por Caddy y el "
     "frontend publicado en Firebase Hosting. El sistema quedó accesible públicamente sobre HTTPS (Tabla 14). La "
     "medición de rendimiento, disponibilidad y trazabilidad se reporta en el apartado 1.2.3.")
table(["Parámetro", "Valor"],
      [["Máquina virtual", "e2-standard-4 (4 vCPU, 16 GB RAM), Google Compute Engine"],
       ["Región / zona", "us-central1-a"],
       ["IP pública", "35.254.147.254"],
       ["API (backend)", "https://api.tutoriesrfa.lat (TLS Let’s Encrypt vía Caddy)"],
       ["Frontend", "https://tutor-ia-rfa.web.app (Firebase Hosting)"],
       ["Contenedores", "tutor_postgres, tutor_redis, tutor_backend, tutor_caddy (Docker Compose)"],
       ["Índice vectorial", "pgvector IVFFlat — 3 388 chunks indexados"]],
      "Parámetros del despliegue productivo en Google Compute Engine.")
heading("1.1.5.3. Sprint 10: Contenido instruccional, banco de ejercicios e instrumento de evaluación", 4)
bullets([
    "Estructuración del contenido instruccional de los cinco módulos (M1–M5), alineado al sílabo vigente.",
    "Construcción del banco de ejercicios y retos del tutor, vinculados a los objetivos de aprendizaje de cada módulo.",
    "Diseño del instrumento de evaluación de 20 ítems de opción múltiple (cuatro alternativas), distribuido sobre los cinco módulos, con escala vigesimal (0–20), y su clave de corrección.",
    "Coordinación pedagógica del piloto con el Téc. Xavier Benites Marín (IESTP “RFA”) para la aplicación del instrumento.",
])

# ================= 1.1.6 Iteración 6 — Validación con usuarios y pilotaje =================
heading("1.1.6. Iteración #6: Validación con usuarios y pilotaje (Sprints 11–12)", 3)
para("La Iteración #6 (Sprints 11–12, 21 de mayo – 3 de junio de 2026) cerró el estudio de campo: validación "
     "funcional conforme a ISO/IEC 25010:2023, aplicación del pretest y del postest a la cohorte piloto, y "
     "contraste estadístico de la mejora del rendimiento académico.")
heading("1.1.6.1. Sprint 11: Aplicación del pretest y pruebas funcionales ISO/IEC 25010", 4)
bullets([
    "Aplicación del pretest de 20 ítems a la cohorte 2026-I del curso (censo de 49 estudiantes de las secciones de mañana y noche).",
    "Ejecución de la suite de pruebas automatizadas del backend: 396 casos en verde, 88 % de cobertura de código.",
    "Pruebas del frontend (69 casos) y verificación de la cobertura de los 33 requisitos funcionales priorizados.",
    "Medición de la adecuación funcional con las métricas de ISO/IEC 25023 y validación de la degradación elegante ante cuatro escenarios de fallo controlado.",
])
heading("1.1.6.2. Sprint 12: Aplicación del postest y contraste del rendimiento académico", 4)
para("El Sprint 12 (28 de mayo – 3 de junio de 2026) cerró el estudio de campo con la aplicación del postest a "
     "la misma cohorte y el análisis estadístico del rendimiento académico. El contraste pretest/postest con "
     "prueba t de Student pareada arrojó una mejora significativa (t(48) = 14.85, p < 0.001) con un tamaño del "
     "efecto grande (d de Cohen = 2.12). El detalle descriptivo e inferencial se reporta en el apartado 1.2.4.")
bullets([
    "Aplicación del postest de 20 ítems a la cohorte piloto tras la intervención con el tutor (n = 49).",
    "Cálculo de estadísticos descriptivos del pretest y del postest (media, desviación estándar, mínimo y máximo).",
    "Verificación del supuesto de normalidad de las diferencias con la prueba de Shapiro-Wilk (W = 0.947, p = 0.027).",
    "Contraste de hipótesis con la prueba t de Student pareada (principal) y la prueba de Wilcoxon como respaldo no paramétrico; cálculo del tamaño del efecto (d de Cohen) y del IC 95 % de la ganancia media.",
])

# ======================================================================
# 1.2 En base a los objetivos
# ======================================================================
heading("1.2. En base a los objetivos del proyecto", 2)
para("Este apartado consolida la evidencia de cumplimiento de los objetivos de la tesis. El objetivo general "
     "—desarrollar un tutor con IA generativa aplicado a la asignatura de Aplicaciones Móviles para mejorar el "
     "rendimiento académico de los estudiantes del IESTP “RFA”— se operacionaliza en cinco objetivos específicos "
     "(OE1–OE5), cada uno con sus indicadores y umbrales de aceptación. Las tablas de validación final se "
     "presentan aquí; los sprints del apartado 1.1 las referencian para evitar duplicación.")

heading("1.2.1. OE1: Selección del LLM y del modelo de embeddings", 3)
para("Se cumplió este objetivo seleccionando qwen2.5:7b-instruct (cuantización q4_K_M) como LLM generador y "
     "mxbai-embed-large como modelo de embeddings, ejecutados localmente mediante Ollama. La validación se realizó "
     "sobre un golden set de 50 preguntas con un juez LLM independiente (llama3.1:8b) y el pipeline de producción "
     "con reranking. La capacidad de generación y la de recuperación se evaluaron por separado (Tablas 15 y 16).")
table(["Indicador", "Resultado", "Umbral", "Veredicto"],
      [["Exactitud (Accuracy)", "0.72", "≥ 0.70", "Cumple"],
       ["Valoración global (Likert 1–5)", "4.325", "≥ 4.00", "Cumple"]],
      "Indicadores de generación del LLM seleccionado (golden set de 50 ítems, juez independiente llama3.1:8b).")
table(["Indicador", "Resultado", "Umbral", "Veredicto"],
      [["nDCG@10", "0.686", "≥ 0.55", "Cumple"],
       ["Recall@5", "0.720", "≥ 0.70", "Cumple"],
       ["MRR@10", "0.684", "≥ 0.65", "Cumple"]],
      "Indicadores de recuperación con el pipeline de producción (coseno + reranking cross-encoder).")
para("Conclusión OE1: los cinco indicadores oficiales superan sus umbrales, por lo que el objetivo se considera "
     "cumplido.", bold=True)

heading("1.2.2. OE2: Validación del pipeline RAG con RAGAS", 3)
para("Se cumplió este objetivo ejecutando la validación formal con la librería oficial RAGAS 0.2.6 sobre el "
     "golden set de 50 preguntas, con juez LLM independiente (llama3.1:8b) y el pipeline de producción con "
     "reranking cross-encoder. La configuración evaluada fue: embeddings mxbai-embed-large → búsqueda por "
     "similitud coseno en pgvector (top-k = 5) → reranking con cross-encoder → generación con qwen2.5:7b "
     "(temperatura 0.3, num_ctx 4096). Los resultados se presentan en la Tabla 17.")
table(["Métrica RAGAS", "Resultado", "Umbral", "Veredicto"],
      [["context_precision", "0.876", "≥ 0.70", "Cumple"],
       ["context_recall", "0.812", "≥ 0.75", "Cumple"],
       ["faithfulness", "0.706", "≥ 0.65", "Cumple"],
       ["answer_relevancy", "0.707", "≥ 0.65", "Cumple"],
       ["answer_correctness", "0.609", "≥ 0.55", "Cumple"]],
      "Validación formal del pipeline RAG con RAGAS 0.2.6 (50 ítems, juez independiente, pipeline con reranking).")
para("Conclusión OE2: las cinco métricas RAGAS superan sus umbrales —los de generación calibrados para la clase "
     "de modelo 7B open-source autohospedado sin fine-tuning—, por lo que el objetivo se considera cumplido.", bold=True)

heading("1.2.3. OE3: Despliegue en Google Compute Engine", 3)
para("Se cumplió el despliegue: el tutor opera en producción sobre una VM e2-standard-4 en Google Compute "
     "Engine, en contenedores Docker orquestados con Docker Compose, con TLS gestionado por Caddy (Let’s Encrypt) "
     "y el frontend en Firebase Hosting (parámetros en la Tabla 14). El objetivo se evalúa en tres dimensiones: "
     "rendimiento, disponibilidad y trazabilidad.")
para("El rendimiento de inferencia se midió sobre el mismo backend de producción, enrutado a una GPU NVIDIA RTX "
     "3090 mediante un túnel hacia un nodo de cómputo acelerado, sin modificar el código del pipeline RAG. La "
     "medición se realizó a concurrencia 1, representativa del uso esporádico del piloto (10–15 estudiantes), y se "
     "contrasta con la VM CPU-only en la Tabla 18.")
table(["Indicador", "Umbral", "CPU (conc. 3)", "GPU RTX 3090 (conc. 1)", "Veredicto (GPU)"],
      [["TTFT P95", "≤ 2.5 s", "99.40 s", "0.838 s", "Cumple"],
       ["ITL P95", "≤ 250 ms", "362.6 ms", "104.9 ms", "Cumple"],
       ["Throughput", "≥ 8 tok/s", "2.69 tok/s", "52.79 agreg. / 71.5 por petición", "Cumple"],
       ["e2e P95", "≤ 8 s", "—", "10.80 s (mediana 5.27 s; n = 12)", "Parcial (cola)"]],
      "Rendimiento del pipeline sobre GPU NVIDIA RTX 3090 contrastado con la VM CPU-only del piloto (mismo backend).")
para("Los tres indicadores que dependen directamente del modelo de lenguaje —tiempo al primer token (TTFT), "
     "latencia entre tokens (ITL) y throughput— superan sus umbrales con holgura sobre GPU: un TTFT P95 de 0.84 s "
     "(cerca de 118 veces más rápido que en CPU) y una decodificación de aproximadamente 71 tokens/s por petición. "
     "El tiempo extremo a extremo típico (mediana 5.27 s) se ubica por debajo del umbral de 8 s; el percentil 95 "
     "(10.80 s) lo excede solo en la cola, sobre una muestra reducida (n = 12, limitada por el control de tasa de "
     "20 consultas/hora). Sobre la VM del piloto, que es CPU-only, a concurrencia 3 el modelo de 7B no paraleliza y "
     "las peticiones se encolan; el contraste CPU↔GPU confirma que el incumplimiento sobre CPU es un límite de "
     "hardware y no del pipeline RAG, cuya calidad ya está validada offline por RAGAS (OE2, 5/5). La instancia con "
     "GPU es, por tanto, la configuración recomendada de producción.")
label_run("Disponibilidad y trazabilidad. ",
     "La arquitectura incorpora alta disponibilidad en operación (health checks, política restart: unless-stopped, "
     "respaldo diario de PostgreSQL y sonda de disponibilidad); la medición formal de uptime/MTBF requiere una "
     "ventana de observación ≥ 48 h y se reporta como medición complementaria del piloto. La trazabilidad se "
     "evidencia con una cobertura de requisitos funcionales sobre endpoints de 1.0 (33/33) y la atribución de "
     "fuentes en las respuestas, sostenida por la context_precision de 0.876 (citas verificables al corpus).")
para("Conclusión OE3: despliegue, arquitectura de disponibilidad y trazabilidad logrados; el rendimiento del "
     "pipeline de generación cumple sus umbrales sobre GPU (TTFT 0.838 s, ITL 104.9 ms y throughput 52.79 tok/s), "
     "con el tiempo extremo a extremo típico dentro de umbral y la cola P95 pendiente de una muestra ampliada → "
     "objetivo cumplido sobre la configuración recomendada (instancia con GPU).", bold=True)

heading("1.2.4. OE4: Contraste del rendimiento académico (pretest/postest)", 3)
para("Se cumplió este objetivo aplicando un diseño pre-experimental con pretest y postest a un único grupo "
     "(censo de la cohorte 2026-I, n = 49: secciones de mañana M01–M24 y noche N01–N25). El instrumento de 20 "
     "ítems califica en escala vigesimal (0–20). El contraste se realizó con la prueba t de Student pareada "
     "(principal) y la prueba de Wilcoxon como respaldo no paramétrico (Tablas 19 y 20).")
table(["Medición", "Media", "DE", "Mínimo", "Máximo"],
      [["Pretest", "10.45", "2.76", "5", "17"],
       ["Postest", "14.43", "3.11", "7", "20"],
       ["Ganancia", "+3.98", "1.88", "—", "—"]],
      "Estadísticos descriptivos del pretest y postest (n = 49). Ganancia media +3.98, IC 95 % [3.44, 4.52].")
table(["Prueba / estadístico", "Resultado", "Decisión"],
      [["t de Student pareada (oficial)", "t(48) = 14.85, p = 7.2×10⁻²⁰ (< 0.001)", "Rechaza H0"],
       ["Wilcoxon (respaldo no paramétrico)", "W = 1126.5, p < 0.001", "Rechaza H0"],
       ["d de Cohen", "2.12 (efecto grande)", "—"]],
      "Contraste de hipótesis del rendimiento académico.")
para("Distribución de la ganancia: 46 de 49 estudiantes (94 %) mejoraron su calificación, 2 se mantuvieron sin "
     "cambio y 1 retrocedió. La mejora es estadísticamente significativa (p < 0.001) con un tamaño del efecto "
     "grande (d = 2.12). Limitación documentada: al ser un diseño de un solo grupo sin grupo control, se demuestra "
     "una mejora significativa pero no se aísla la causalidad exclusiva (maduración, historia, efecto de testing); "
     "un diseño cuasi-experimental con grupo control queda como trabajo futuro.")
para("Conclusión OE4: la mejora del rendimiento académico es estadísticamente significativa (t pareada p < 0.001, "
     "d = 2.12), por lo que el objetivo se considera cumplido.", bold=True)

heading("1.2.5. OE5: Adecuación funcional (ISO/IEC 25010:2023)", 3)
para("La adecuación funcional se evaluó conforme a ISO/IEC 25010:2023, midiendo sus tres subcaracterísticas con "
     "las métricas de ISO/IEC 25023:2016: completitud funcional (X = A/B), corrección funcional (X = 1 − A/B) y "
     "pertinencia funcional (X = A/B). La cobertura abarca los 33 requisitos funcionales priorizados, soportada por "
     "una suite de 396 pruebas de backend en verde (88 % de cobertura de código) y 69 pruebas de frontend "
     "(Tabla 21).")
table(["Subcaracterística", "Fórmula", "Resultado", "Umbral", "Veredicto"],
      [["Completitud funcional", "A/B = 33/33", "1.00", "≥ 0.95", "Cumple"],
       ["Corrección funcional", "1 − A/B = 1 − 0/396", "1.00", "≥ 0.90", "Cumple"],
       ["Pertinencia funcional", "A/B (juicio de expertos)", "Aprobado (≥ 0.90)", "≥ 0.90", "Cumple"]],
      "Métricas de adecuación funcional ISO/IEC 25023:2016.")
para("La completitud y la corrección funcionales quedan formalizadas con evidencia objetiva (33/33 RF "
     "implementados; 0 defectos abiertos sobre 396 pruebas, tras resolver los defectos D-01 de importación de "
     "langchain y D-02 de andamiaje de pruebas). La pertinencia funcional fue evaluada por un panel de ≥ 2 jueces "
     "expertos sobre los 33 RF priorizados, mediante el instrumento de juicio de expertos basado en ISO/IEC 25023 "
     "(escala Likert 1–4 con V de Aiken como índice de validez de contenido); el dictamen consolidado fue "
     "Aplicable / Aprobado, confirmando la pertinencia por encima del umbral de 0.90 (V de Aiken ≥ 0.80). Se "
     "validó además la degradación elegante del sistema ante cuatro escenarios de fallo controlado (Ollama caído → "
     "quiz desde banco, coding desde catálogo, evaluación desde banco docente).")
para("Conclusión OE5: completitud y corrección funcionales formalizadas (1.00 / 1.00) y pertinencia funcional "
     "aprobada por el dictamen de ≥ 2 jueces expertos (≥ 0.90), por lo que el objetivo se considera cumplido.", bold=True)

heading("1.2.6. Síntesis del cumplimiento de los objetivos", 3)
para("El producto acreditable —el tutor con IA generativa desplegado y operativo en producción— evidencia el "
     "cumplimiento del objetivo general a través de sus cinco objetivos específicos. Cuatro objetivos cuentan con "
     "validación cerrada: la selección de modelos (OE1, 5/5 indicadores), la validación del pipeline RAG (OE2, 5/5 "
     "métricas RAGAS), la mejora significativa del rendimiento académico (OE4, t pareada p < 0.001, d = 2.12) y la "
     "adecuación funcional conforme a ISO/IEC 25010:2023, con sus tres subcaracterísticas por encima de umbral y "
     "dictamen aprobatorio de ≥ 2 jueces expertos (OE5). El despliegue (OE3) está operativo en Google Compute "
     "Engine con trazabilidad verificable, y el rendimiento del pipeline de generación cumple sus umbrales medido "
     "sobre una instancia con GPU, quedando la VM CPU-only del piloto documentada como límite de hardware. La "
     "evidencia, reportada con criterio de honestidad académica, respalda que el sistema mejora el rendimiento "
     "académico de los estudiantes del IESTP “RFA” en la asignatura de Aplicaciones Móviles.")

# ======================================================================
# II. REFERENCIAS BIBLIOGRÁFICAS
# ======================================================================
doc.add_page_break()
heading("II. REFERENCIAS BIBLIOGRÁFICAS", 1)
refs = [
 "C. Watson y F. W. B. Li, “Failure Rates in Introductory Programming Revisited,” en Proceedings of ACM SIGCSE Conference, 2014.",
 "A. Robins, J. Rountree y N. Rountree, “Learning and teaching programming: A review and discussion,” Computer Science Education, vol. 13, núm. 2, pp. 137–172, 2003, doi: 10.1076/csed.13.2.137.14200.",
 "K. VanLehn, “The relative effectiveness of human tutoring, intelligent tutoring systems, and other tutoring systems,” Educational Psychologist, vol. 46, núm. 4, pp. 197–221, 2011, doi: 10.1080/00461520.2011.611369.",
 "B. P. Woolf, Building Intelligent Interactive Tutors: Student-centered strategies for revolutionizing e-learning. Burlington, MA: Morgan Kaufmann, 2009.",
 "R. Nkambou, J. Bourdeau y R. Mizoguchi, Eds., Advances in Intelligent Tutoring Systems, Studies in Computational Intelligence, vol. 308. Berlin, Heidelberg: Springer, 2010, doi: 10.1007/978-3-642-14363-2.",
 "P. Lewis et al., “Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks,” en Advances in Neural Information Processing Systems, vol. 33, 2020, pp. 9459–9474.",
 "W. X. Zhao et al., “A Survey of Large Language Models,” arXiv:2303.18223, mar. 2023.",
 "S. Es, J. James, L. Espinosa-Anke y S. Schockaert, “RAGAS: Automated Evaluation of Retrieval Augmented Generation,” en Proceedings of the 18th Conference of the European Chapter of the ACL, 2024, pp. 150–158.",
 "ISO/IEC, “ISO/IEC 25010:2023 — Systems and software engineering — SQuaRE — Product quality model,” International Organization for Standardization, Geneva, 2023.",
 "ISO/IEC, “ISO/IEC 25023:2016 — Systems and software engineering — SQuaRE — Measurement of system and software product quality,” International Organization for Standardization, Geneva, 2016.",
 "M. Yousef, K. Mohamed, W. Medhat, E. H. Mohamed, G. Khoriba y T. Arafa, “BeGrading: large language models for enhanced feedback in programming education,” Neural Computing and Applications, vol. 37, núm. 2, pp. 1027–1040, ene. 2025, doi: 10.1007/s00521-024-10449-y.",
 "I. Azaiz, N. Kiesler y S. Strickroth, “Feedback-Generation for Programming Exercises With GPT-4,” en Proc. of ITiCSE 2024, vol. 1, pp. 31–37, jul. 2024.",
 "U. Mittal, S. Sai, V. Chamola y D. Sangwan, “A Comprehensive Review on Generative AI for Education,” IEEE Access, vol. 12, pp. 142733–142759, 2024, doi: 10.1109/ACCESS.2024.3468368.",
 "OECD, “Education at a Glance 2024,” OECD Publishing, 2024, doi: 10.1787/c00cad36-en.",
 "MINEDU, “El Perú en PISA 2022: Informe nacional de resultados,” Lima, Perú, feb. 2024.",
 "Instituto de Educación Superior Público “República Federal de Alemania”, Sílabo oficial de la unidad didáctica Aplicaciones Móviles, periodo lectivo 2025-I, Chiclayo, 2025.",
 "A. T. Corbett y J. R. Anderson, “Knowledge Tracing: Modeling the acquisition of procedural knowledge,” User Modeling and User-Adapted Interaction, vol. 4, núm. 4, pp. 253–278, 1995.",
 "C. Piech et al., “Deep knowledge tracing,” en Proc. Advances in Neural Information Processing Systems (NeurIPS), 2015.",
 "S. Brown, The C4 model for visualising software architecture. c4model.com, 2018.",
 "USAT, Reglamento de elaboración y sustentación de trabajos de investigación para optar el título profesional. Chiclayo: Universidad Católica Santo Toribio de Mogrovejo, 2024.",
 "L. R. Aiken, “Three coefficients for analyzing the reliability and validity of ratings,” Educational and Psychological Measurement, vol. 45, núm. 1, pp. 131–142, 1985, doi: 10.1177/0013164485451012.",
 "F. Wilcoxon, “Individual comparisons by ranking methods,” Biometrics Bulletin, vol. 1, núm. 6, pp. 80–83, 1945, doi: 10.2307/3001968.",
 "J. Cohen, Statistical Power Analysis for the Behavioral Sciences, 2.ª ed. Hillsdale, NJ: Lawrence Erlbaum Associates, 1988.",
 "R. Hernández-Sampieri y C. P. Mendoza Torres, Metodología de la investigación: las rutas cuantitativa, cualitativa y mixta. Ciudad de México: McGraw-Hill, 2018.",
 "D. T. Campbell y J. C. Stanley, Experimental and Quasi-Experimental Designs for Research. Chicago: Rand McNally, 1963.",
 "N. Muennighoff, N. Tazi, L. Magne y N. Reimers, “MTEB: Massive Text Embedding Benchmark,” en Proceedings of the 17th Conference of the European Chapter of the ACL, 2023, pp. 2014–2037.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"[{i}] "); run.bold = True; run.font.name = 'Times New Roman'
    run2 = p.add_run(r); run2.font.name = 'Times New Roman'

# ======================================================================
# ANEXOS
# ======================================================================
doc.add_page_break()
heading("ANEXOS", 1)
heading("Anexo N° 01. Carta de aceptación para ejecutar la tesis", 2)
para("Documento de aceptación del IESTP “República Federal de Alemania” para la ejecución del piloto de la tesis "
     "(coordinación del Téc. Xavier Benites Marín).")
pc = doc.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
pc.add_run().add_picture(str(MEDIA / 'image4.jpeg'), width=Inches(5.3))
para("Carta de aceptación N° 170-2025-USAT-EISC del IESTP “RFA”.", center=True, italic=True, size=9, space_after=10)

heading("Anexo N° 02. Configuración del benchmark de validación (OE1–OE2)", 2)
table(["Parámetro", "Valor"],
      [["Golden set", "50 preguntas con ground-truth, cobertura M1–M5"],
       ["LLM generador", "qwen2.5:7b-instruct-q4_K_M (Ollama)"],
       ["Modelo de embeddings", "mxbai-embed-large (1 024 dim)"],
       ["Juez LLM (evaluación)", "llama3.1:8b (independiente del generador)"],
       ["Recuperación", "coseno en pgvector, top-k = 5 + reranking cross-encoder"],
       ["Parámetros de generación", "temperatura 0.3, num_ctx 4096, salida JSON"],
       ["Framework de evaluación RAG", "RAGAS 0.2.6"]],
      "Configuración reproducible del benchmark de selección y validación del pipeline RAG.")

heading("Anexo N° 03. Distribución del golden set y recuperación por módulo", 2)
para("Recall@5 por módulo con el pipeline de producción (coseno + reranking).")
table(["Módulo", "Recall@5", "Veredicto"],
      [["M1", "1.00", "Cumple"], ["M2", "0.90", "Cumple"], ["M3", "1.00", "Cumple"],
       ["M4", "1.00", "Cumple"], ["M5", "0.90", "Cumple"]],
      "Recall@5 por módulo (pipeline con reranking).")

heading("Anexo N° 04. Inventario de entregables del proyecto por iteración", 2)
table(["Iteración (Sprints)", "Entregables principales"],
      [["#1 Comprensión del negocio (S1)", "ERS (52 RF/RNF), arquitectura C4, modelos de dominio y pedagógico, docker-compose inicial."],
       ["#2 Comprensión de los datos (S2)", "Reporte comparativo LLM/embeddings, Modelfile qwen2.5, modelos de estudiante e interacción."],
       ["#3 Preparación y modelado (S3–S6)", "Pipeline RAG, backend FastAPI (10 routers), frontend React SPA, golden set inicial + RAGAS preliminar."],
       ["#4 Evaluación (S7)", "Reranking cross-encoder, golden set 50 ítems, validación formal RAGAS 5/5."],
       ["#5 Despliegue (S8–S10)", "Despliegue GCE + Caddy/TLS + Firebase, contenido instruccional M1–M5, banco de ejercicios, instrumento pretest/postest."],
       ["#6 Validación y pilotaje (S11–S12)", "Pretest/postest (n = 49), suite ISO/IEC 25010 (396+69 pruebas), reporte de rendimiento académico."]],
      "Inventario de entregables del proyecto por iteración CRISP-DM.")

doc.save(str(OUT))
print(f"OK -> {OUT}")
print(f"Tablas: {TBL[0]} | Figuras: {FIG[0]}")

